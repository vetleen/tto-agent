"""Organization "Styles" — design levers for exported documents.

Single source of truth for the small set of typography/colour levers an
organization can configure (Org settings → Styles) and that the canvas → Word
export applies. Kept provider/feature-agnostic so other layouts (e.g. a future
canvas preview) can reuse the same values.

Stored on ``Organization.preferences["styles"]`` (no dedicated model field).
The web settings path uses :func:`get_org_styles`, :func:`validate_styles` and
:data:`FONT_CHOICES`; the export path additionally uses :func:`apply_doc_styles`
(which imports python-docx lazily so the settings page stays light).
"""
from __future__ import annotations

import base64
import io
import logging
import re
from dataclasses import dataclass

logger = logging.getLogger(__name__)

# Zero-width space that chat.services.render_citations injects at the start of
# each numbered source item, so the export can find and shrink just that list
# (a normal numbered list is untouched). Stripped during export.
FOOTNOTE_MARKER = "​"

# Footnote/source list and table text each render this many points below body.
FOOTNOTE_SIZE_DELTA = 2
TABLE_SIZE_DELTA = 2

# --- Fonts ---------------------------------------------------------------
#
# A ``.docx`` stores only a font *name*; Word substitutes locally if it isn't
# installed, so DOCX export is name-only. PDF export (WeasyPrint) instead needs
# the actual font *file* on the server, resolved by ``core.fonts``. To keep PDF
# output faithful we bundle OFL/Apache fonts (``core/assets/fonts/<key>/``) and
# map common proprietary names onto them.
#
# BUNDLED_FONTS: families we ship (four faces each: Regular/Bold/Italic/
# BoldItalic). ``category`` drives the CSS generic fallback.
BUNDLED_FONTS = {
    "Carlito": {"label": "Carlito", "category": "sans"},
    "Arimo": {"label": "Arimo", "category": "sans"},
    "Caladea": {"label": "Caladea", "category": "serif"},
    "Tinos": {"label": "Tinos", "category": "serif"},
    "Gelasio": {"label": "Gelasio", "category": "serif"},
    "EBGaramond": {"label": "EB Garamond", "category": "serif"},
    "Cousine": {"label": "Cousine", "category": "mono"},
}

# FONT_SUBSTITUTES: normalised (lowercased) proprietary/common name ->
# (bundled family key, fidelity). "metric" = metric-compatible, visually faithful
# (silent in PDF export); "visual" = close but not metric-identical (surfaces a
# soft note). Keys are matched after :func:`core.fonts.normalize_font_name`.
FONT_SUBSTITUTES = {
    "calibri": ("Carlito", "metric"),
    "arial": ("Arimo", "metric"),
    "helvetica": ("Arimo", "visual"),
    "cambria": ("Caladea", "metric"),
    "times new roman": ("Tinos", "metric"),
    "times": ("Tinos", "visual"),
    "georgia": ("Gelasio", "metric"),
    "courier new": ("Cousine", "metric"),
    "courier": ("Cousine", "visual"),
    "garamond": ("EBGaramond", "visual"),
    "segoe ui": ("Carlito", "visual"),
    "verdana": ("Arimo", "visual"),
    "tahoma": ("Arimo", "visual"),
    "trebuchet ms": ("Arimo", "visual"),
    "book antiqua": ("Gelasio", "visual"),
    "palatino": ("Gelasio", "visual"),
}

# Neutral fallback when a requested font resolves nowhere (bundled, substitute,
# uploaded, or Google Fonts). A clean professional sans.
FALLBACK_FONT = "Carlito"

# Quick-pick groups for the settings picker. Every name here is "solvable": we
# either ship it (BUNDLED_FONTS) or substitute it (FONT_SUBSTITUTES). The UI also
# offers a "Custom…" option (sentinel below) for any name the admin types — those
# resolve via Google Fonts at export time, or fall back with a warning.
FONT_CHOICES = [
    {"group": "Sans-serif", "fonts": ["Calibri", "Arial", "Segoe UI", "Verdana",
                                       "Tahoma", "Trebuchet MS", "Carlito", "Arimo"]},
    {"group": "Serif", "fonts": ["Cambria", "Times New Roman", "Georgia", "Garamond",
                                 "Book Antiqua", "Caladea", "Tinos", "Gelasio", "EB Garamond"]},
    {"group": "Monospace", "fonts": ["Courier New", "Cousine"]},
]
CUSTOM_FONT_SENTINEL = "__custom__"

# Defaults match today's html2docx output closely (Calibri 11, near-black text)
# so exports for orgs that never touch the setting are essentially unchanged.
# accent_color defaults to "" (empty = opt-in): table headers are only shaded
# once an org sets it.
STYLE_DEFAULTS = {
    "body_font": "Calibri",
    "body_size": 11,
    "heading_font": "Calibri",
    "heading_color": "#1A1A1A",
    "body_color": "#1A1A1A",
    "accent_color": "",
    # Tables. Header rows reuse accent_color for shading; these add borders and
    # optional zebra striping. bold-header, repeat-on-page-break and cell padding
    # are applied unconditionally (sensible defaults, no setting).
    "table_border_style": "all",      # all | horizontal | header | none
    "table_border_color": "#CCCCCC",
    "table_banded": False,
    # Page header & footer (independent typography each). Empty text means "no
    # text"; the page number ("3 / 12") sits bottom-right of the footer.
    "header_text": "",
    "header_font": "Calibri",
    "header_size": 9,
    "header_color": "#1A1A1A",
    "header_bold": False,
    "header_italic": False,
    # Optional brand logo in the header. "none" = no logo; "left"/"right" place
    # the logo (Organization.logo) on that side, pushing header_text to the
    # other. Footer carries no logo.
    "header_logo_position": "none",
    "footer_text": "",
    "footer_font": "Calibri",
    "footer_size": 9,
    "footer_color": "#1A1A1A",
    "footer_bold": False,
    "footer_italic": False,
    "footer_page_numbers": True,
}

TABLE_BORDER_STYLES = ("all", "horizontal", "header", "none")
HEADER_LOGO_POSITIONS = ("none", "left", "right")
HEADER_FOOTER_TEXT_MAX = 200

# Logo print box in the document header. The logo is scaled proportionally to
# fit within these (DOCX picture size / PDF background-size), so neither
# dimension is ever exceeded and the aspect ratio is preserved. Height stays
# well under the 2.54cm page margin so it never collides with the body.
LOGO_PRINT_MAX_H_CM = 1.4
LOGO_PRINT_MAX_W_CM = 6.0

_FONT_RE = re.compile(r"^[A-Za-z0-9 \-]+$")  # curated + custom like "PT Sans"
_HEX_RE = re.compile(r"^#[0-9A-Fa-f]{6}$")


def get_org_styles(org) -> dict:
    """Resolve an org's effective styles (defaults overlaid with stored values).

    ``org`` may be ``None`` (no membership) → defaults. Pure dict read, no DB,
    so it's safe to call from the async export view. Stored values are written
    through :func:`validate_styles`, so they're trusted here.
    """
    styles = dict(STYLE_DEFAULTS)
    stored = (getattr(org, "preferences", None) or {}).get("styles") if org is not None else None
    if isinstance(stored, dict):
        for key in STYLE_DEFAULTS:
            if key in stored and stored[key] is not None:
                styles[key] = stored[key]
    return styles


def validate_styles(data) -> tuple[dict | None, str | None]:
    """Validate a styles payload from the settings endpoint.

    Returns ``(clean_dict, None)`` on success or ``(None, error_message)``.
    Missing keys fall back to :data:`STYLE_DEFAULTS`.
    """
    if not isinstance(data, dict):
        return None, "Invalid styles payload."

    clean = dict(STYLE_DEFAULTS)

    for key in ("body_font", "heading_font"):
        val = data.get(key, clean[key])
        if not isinstance(val, str):
            return None, f"{key.replace('_', ' ').title()} must be text."
        val = val.strip()
        if not val or len(val) > 64 or not _FONT_RE.match(val):
            return None, "Font names may use letters, numbers, spaces and hyphens (max 64 chars)."
        clean[key] = val

    size = data.get("body_size", clean["body_size"])
    try:
        size = int(size)
    except (TypeError, ValueError):
        return None, "Body size must be a number."
    if not (8 <= size <= 24):
        return None, "Body size must be between 8 and 24."
    clean["body_size"] = size

    for key in ("heading_color", "body_color"):
        val = data.get(key, clean[key])
        if not isinstance(val, str) or not _HEX_RE.match(val.strip()):
            return None, f"{key.replace('_', ' ').title()} must be a hex colour like #1A1A1A."
        clean[key] = val.strip().upper()

    # Accent is optional: "" means "no table-header shading".
    accent = data.get("accent_color", clean["accent_color"])
    if not isinstance(accent, str):
        return None, "Accent colour must be text."
    accent = accent.strip()
    if accent and not _HEX_RE.match(accent):
        return None, "Accent colour must be a hex colour like #2563EB, or left blank."
    clean["accent_color"] = accent.upper() if accent else ""

    border_style = data.get("table_border_style", clean["table_border_style"])
    if border_style not in TABLE_BORDER_STYLES:
        return None, "Invalid table border style."
    clean["table_border_style"] = border_style

    bcol = data.get("table_border_color", clean["table_border_color"])
    if not isinstance(bcol, str) or not _HEX_RE.match(bcol.strip()):
        return None, "Table border colour must be a hex colour like #CCCCCC."
    clean["table_border_color"] = bcol.strip().upper()

    banded = data.get("table_banded", clean["table_banded"])
    if not isinstance(banded, bool):
        return None, "Banded rows must be true or false."
    clean["table_banded"] = banded

    for prefix in ("header", "footer"):
        text = data.get(f"{prefix}_text", clean[f"{prefix}_text"])
        if not isinstance(text, str):
            return None, f"{prefix.title()} text must be text."
        text = text.strip()
        if len(text) > HEADER_FOOTER_TEXT_MAX:
            return None, f"{prefix.title()} text must be {HEADER_FOOTER_TEXT_MAX} characters or fewer."
        clean[f"{prefix}_text"] = text

        font = data.get(f"{prefix}_font", clean[f"{prefix}_font"])
        font = font.strip() if isinstance(font, str) else ""
        if not font or len(font) > 64 or not _FONT_RE.match(font):
            return None, f"{prefix.title()} font names may use letters, numbers, spaces and hyphens (max 64 chars)."
        clean[f"{prefix}_font"] = font

        size = data.get(f"{prefix}_size", clean[f"{prefix}_size"])
        try:
            size = int(size)
        except (TypeError, ValueError):
            return None, f"{prefix.title()} size must be a number."
        if not (6 <= size <= 24):
            return None, f"{prefix.title()} size must be between 6 and 24."
        clean[f"{prefix}_size"] = size

        color = data.get(f"{prefix}_color", clean[f"{prefix}_color"])
        if not isinstance(color, str) or not _HEX_RE.match(color.strip()):
            return None, f"{prefix.title()} colour must be a hex colour like #1A1A1A."
        clean[f"{prefix}_color"] = color.strip().upper()

        for suffix in ("bold", "italic"):
            val = data.get(f"{prefix}_{suffix}", clean[f"{prefix}_{suffix}"])
            if not isinstance(val, bool):
                return None, f"{prefix.title()} {suffix} must be true or false."
            clean[f"{prefix}_{suffix}"] = val

    page_numbers = data.get("footer_page_numbers", clean["footer_page_numbers"])
    if not isinstance(page_numbers, bool):
        return None, "Page numbers must be true or false."
    clean["footer_page_numbers"] = page_numbers

    logo_position = data.get("header_logo_position", clean["header_logo_position"])
    if logo_position not in HEADER_LOGO_POSITIONS:
        return None, "Invalid logo position."
    clean["header_logo_position"] = logo_position

    return clean, None


# --- Header logo ------------------------------------------------------------

@dataclass
class LogoRender:
    """A resolved header logo, ready for both export paths.

    ``data`` is the (PNG) bytes read once from storage; ``aspect`` is width/height;
    ``position`` is "left" or "right". Built by :func:`resolve_org_logo` at export
    time and threaded into :func:`apply_doc_styles` / :func:`build_pdf_css`.
    """
    data: bytes
    aspect: float
    position: str  # "left" | "right"


def _logo_print_dims(aspect: float) -> tuple[float, float]:
    """Logo ``(width_cm, height_cm)`` fitting the print box, aspect preserved.

    Caps height first, then width — so neither ``LOGO_PRINT_MAX_*`` is exceeded
    regardless of how wide or tall the source is. Both DOCX (picture size) and PDF
    (``background-size``) use this so the two formats render the logo identically.
    """
    if aspect <= 0:
        return LOGO_PRINT_MAX_W_CM, LOGO_PRINT_MAX_H_CM
    height = LOGO_PRINT_MAX_H_CM
    width = height * aspect
    if width > LOGO_PRINT_MAX_W_CM:
        width = LOGO_PRINT_MAX_W_CM
        height = width / aspect
    return width, height


def resolve_org_logo(styles: dict, org) -> LogoRender | None:
    """Read the org's header logo for export, or ``None`` when there's none.

    Returns ``None`` unless ``header_logo_position`` is "left"/"right" *and* the
    org has a stored logo. Reads the blob (I/O — call from a sync context) plus the
    image's intrinsic size. Never raises: any failure logs and returns ``None`` so
    a broken or missing logo can't take down an export.
    """
    position = (styles.get("header_logo_position") or "none").strip()
    if position not in ("left", "right"):
        return None
    logo = getattr(org, "logo", None)
    if not logo:
        return None
    try:
        from PIL import Image

        with logo.open("rb") as fh:
            data = fh.read()
        with Image.open(io.BytesIO(data)) as img:
            width, height = img.size
        if not data or not width or not height:
            return None
        return LogoRender(data=data, aspect=width / height, position=position)
    except Exception:  # noqa: BLE001
        logger.warning("Org logo resolution failed", exc_info=True)
        return None


def _norm_hex(value) -> str | None:
    """Return a 6-char uppercase hex string (no ``#``) or ``None``."""
    if not isinstance(value, str):
        return None
    v = value.strip().lstrip("#")
    if len(v) == 6 and all(c in "0123456789abcdefABCDEF" for c in v):
        return v.upper()
    return None


def _readable_text(hex6: str) -> str:
    """Pick black or white text for a given background (perceived luminance)."""
    r, g, b = int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16)
    luminance = 0.299 * r + 0.587 * g + 0.114 * b
    return "1A1A1A" if luminance > 150 else "FFFFFF"


def _set_style_font(style, font_name: str) -> None:
    """Force a paragraph style's font name to ``font_name``.

    Word's built-in styles reference a *theme* font (e.g.
    ``w:rFonts w:asciiTheme="majorHAnsi"``), which Word prefers over an explicit
    ``w:ascii`` when both are present — so simply setting ``style.font.name``
    is silently ignored for headings. We set the explicit Latin font names and
    drop the theme references so the chosen font actually renders. Complex-script
    and East-Asian theme refs are left intact (they govern non-Latin scripts).
    """
    from docx.oxml.ns import qn

    rfonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    for theme_attr in ("w:asciiTheme", "w:hAnsiTheme"):
        rfonts.attrib.pop(qn(theme_attr), None)


def _tint(hex6: str, fraction: float) -> str:
    """Blend a colour toward white by ``fraction`` (0 = colour, 1 = white)."""
    r, g, b = int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16)
    r = int(r + (255 - r) * fraction)
    g = int(g + (255 - g) * fraction)
    b = int(b + (255 - b) * fraction)
    return f"{r:02X}{g:02X}{b:02X}"


def _style_tables(doc, styles: dict) -> None:
    """Apply border/shading/banding to every table in the document.

    Header rows (row 0) get bold text, optional accent shading (with auto
    contrast text), and repeat on page breaks. Borders, cell padding and
    optional zebra striping follow the org settings. OOXML elements are inserted
    in schema order so Word doesn't flag the file for repair.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    accent = _norm_hex(styles.get("accent_color"))
    border_style = styles.get("table_border_style") or "all"
    border_color = _norm_hex(styles.get("table_border_color")) or "CCCCCC"
    banded = bool(styles.get("table_banded"))
    band_hex = _tint(accent, 0.85) if accent else "F2F2F2"
    try:
        body_size = int(styles.get("body_size"))
    except (TypeError, ValueError):
        body_size = STYLE_DEFAULTS["body_size"]
    table_size = Pt(max(1, body_size - TABLE_SIZE_DELTA))

    def make(tag, **attrs):
        el = OxmlElement(f"w:{tag}")
        for key, value in attrs.items():
            el.set(qn(f"w:{key}"), value)
        return el

    # Successor tags (schema order) used to position inserted elements.
    # insert_element_before() qualifies these names itself — pass them raw.
    _TCPR_AFTER_BORDERS = ("w:shd", "w:noWrap", "w:tcMar", "w:textDirection",
                           "w:tcFitText", "w:vAlign", "w:hideMark")
    _TCPR_AFTER_SHD = _TCPR_AFTER_BORDERS[1:]

    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        old = tcPr.find(qn("w:shd"))
        if old is not None:
            tcPr.remove(old)
        tcPr.insert_element_before(make("shd", val="clear", color="auto", fill=fill), *_TCPR_AFTER_SHD)

    def header_rule(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is None:
            borders = make("tcBorders")
            tcPr.insert_element_before(borders, *_TCPR_AFTER_BORDERS)
        old = borders.find(qn("w:bottom"))
        if old is not None:
            borders.remove(old)
        borders.append(make("bottom", val="single", sz="8", space="0", color=border_color))

    def table_borders(table):
        tblPr = table._tbl.tblPr
        old = tblPr.find(qn("w:tblBorders"))
        if old is not None:
            tblPr.remove(old)
        if border_style not in ("all", "horizontal"):
            return
        edges = (["top", "left", "bottom", "right", "insideH", "insideV"]
                 if border_style == "all" else ["top", "bottom", "insideH"])
        borders = make("tblBorders")
        for edge in edges:
            borders.append(make(edge, val="single", sz="4", space="0", color=border_color))
        tblPr.insert_element_before(borders, "w:shd", "w:tblLayout",
                                    "w:tblCellMar", "w:tblLook")

    def cell_margins(table):
        tblPr = table._tbl.tblPr
        old = tblPr.find(qn("w:tblCellMar"))
        if old is not None:
            tblPr.remove(old)
        mar = make("tblCellMar")
        for side in ("top", "left", "bottom", "right"):
            mar.append(make(side, w="80", type="dxa"))  # 80 dxa ≈ 0.06"
        tblPr.insert_element_before(mar, "w:tblLook")

    def repeat_header(row):
        trPr = row._tr.get_or_add_trPr()
        if trPr.find(qn("w:tblHeader")) is None:
            trPr.append(make("tblHeader", val="true"))

    for table in doc.tables:
        table_borders(table)
        cell_margins(table)
        if not table.rows:
            continue

        # Table text renders a couple of points below body, like the sources list.
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = table_size

        header = table.rows[0]
        repeat_header(header)
        header_text_hex = _readable_text(accent) if accent else None
        for cell in header.cells:
            if accent:
                shade(cell, accent)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    if header_text_hex:
                        run.font.color.rgb = RGBColor.from_string(header_text_hex)
        if border_style == "header":
            for cell in header.cells:
                header_rule(cell)

        if banded:
            for body_index, row in enumerate(table.rows[1:]):
                if body_index % 2 == 1:  # 2nd, 4th, … body row
                    for cell in row.cells:
                        shade(cell, band_hex)


def _resize_footnote_sources(doc, size_pt: int) -> None:
    """Shrink the footnote/source list to ``size_pt``.

    Targets only paragraphs whose first run begins with :data:`FOOTNOTE_MARKER`
    (inserted by ``render_citations``). Strips the marker, sizes every run, and
    sets the paragraph-mark run properties so the auto-generated list number
    matches the text.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    half_points = str(int(size_pt * 2))
    for para in doc.paragraphs:
        if not para.runs or not para.runs[0].text.startswith(FOOTNOTE_MARKER):
            continue
        para.runs[0].text = para.runs[0].text.replace(FOOTNOTE_MARKER, "", 1)
        for run in para.runs:
            run.font.size = Pt(size_pt)
        pPr = para._p.get_or_add_pPr()
        rPr = pPr.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            pPr.append(rPr)  # mark rPr is last in the pPr sequence
        old = rPr.find(qn("w:sz"))
        if old is not None:
            rPr.remove(old)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), half_points)
        rPr.append(sz)


def _add_page_field(paragraph, instr: str) -> None:
    """Append a simple field (e.g. ``PAGE`` / ``NUMPAGES``) to a paragraph.

    python-docx has no field API, so we drop in ``<w:fldSimple>`` with a cached
    "1" the reader recomputes. Inherits the paragraph's (Footer) style.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" {instr} ")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def _style_hf(doc, style_name: str, styles: dict, prefix: str) -> None:
    """Apply one of the ``Header``/``Footer`` styles from ``{prefix}_*`` keys."""
    from docx.shared import Pt, RGBColor

    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    _set_style_font(style, styles.get(f"{prefix}_font") or STYLE_DEFAULTS[f"{prefix}_font"])
    try:
        style.font.size = Pt(int(styles.get(f"{prefix}_size")))
    except (TypeError, ValueError):
        style.font.size = Pt(STYLE_DEFAULTS[f"{prefix}_size"])
    color = _norm_hex(styles.get(f"{prefix}_color"))
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bool(styles.get(f"{prefix}_bold"))
    style.font.italic = bool(styles.get(f"{prefix}_italic"))


def _add_header_logo(doc, section, para, logo, header_text: str) -> None:
    """Lay out the header logo on its side with the text on the other.

    Uses a single right-aligned tab stop at the content width (the same trick the
    footer uses for page numbers): content before the tab stays left, content
    after is flush right. So logo-left ⇒ ``[logo]\\t[text]`` and logo-right ⇒
    ``[text]\\t[logo]``. Sizing comes from :func:`_logo_print_dims`, so the
    picture matches the PDF logo exactly.
    """
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Cm

    width_cm, height_cm = _logo_print_dims(logo.aspect)

    def add_picture():
        para.add_run().add_picture(io.BytesIO(logo.data), width=Cm(width_cm), height=Cm(height_cm))

    need_tab = (logo.position == "left" and header_text) or logo.position == "right"
    if need_tab:
        # Clear the built-in Header style's center/right tabs (sized for 1"
        # margins) and add one right tab at this doc's content width, mirroring
        # how the footer positions the page number.
        content_width = section.page_width - section.left_margin - section.right_margin
        try:
            tab_stops = doc.styles["Header"].paragraph_format.tab_stops
            tab_stops.clear_all()
            tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
        except KeyError:
            para.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)

    if logo.position == "left":
        add_picture()
        if header_text:
            para.add_run("\t")
            para.add_run(header_text)
    else:  # right
        if header_text:
            para.add_run(header_text)
        para.add_run("\t")
        add_picture()


def _apply_header_footer(doc, styles: dict, logo=None) -> None:
    """Add the org's page header/footer with independent typography.

    Header and footer each take their own font/size/colour/bold/italic, applied
    via the built-in ``Header``/``Footer`` styles so every run — including the
    page-number field — inherits the formatting. An optional ``logo``
    (:class:`LogoRender`) sits on its configured side of the header with the text
    on the other. The footer lays out ``footer_text`` on the left and a ``3 / 12``
    page indicator on the right via a right-aligned tab stop.
    """
    from docx.enum.text import WD_TAB_ALIGNMENT

    header_text = (styles.get("header_text") or "").strip()
    footer_text = (styles.get("footer_text") or "").strip()
    page_numbers = bool(styles.get("footer_page_numbers"))
    has_logo = logo is not None
    if not header_text and not footer_text and not page_numbers and not has_logo:
        return

    if header_text or has_logo:
        _style_hf(doc, "Header", styles, "header")
    if footer_text or page_numbers:
        _style_hf(doc, "Footer", styles, "footer")

    section = doc.sections[0]

    if header_text or has_logo:
        section.header.is_linked_to_previous = False
        para = section.header.paragraphs[0]
        para.text = ""
        if has_logo:
            _add_header_logo(doc, section, para, logo, header_text)
        else:
            para.text = header_text

    if footer_text or page_numbers:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0]
        para.text = ""
        if page_numbers:
            # The built-in Footer style ships with center + right tabs sized for
            # 1" margins, so a single \t lands the page number at center. Replace
            # them with one right tab at this doc's content width so it sits flush
            # right regardless of margins.
            content_width = section.page_width - section.left_margin - section.right_margin
            try:
                tab_stops = doc.styles["Footer"].paragraph_format.tab_stops
                tab_stops.clear_all()
                tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
            except KeyError:
                para.paragraph_format.tab_stops.add_tab_stop(content_width, WD_TAB_ALIGNMENT.RIGHT)
        if footer_text:
            para.add_run(footer_text)
        if page_numbers:
            para.add_run("\t")
            _add_page_field(para, "PAGE")
            para.add_run(" / ")
            _add_page_field(para, "NUMPAGES")


def apply_doc_styles(doc, styles: dict, logo=None) -> None:
    """Apply resolved org styles to a python-docx ``Document`` in place.

    Sets the ``Normal`` style (body font/size/colour) and ``Heading 1``–``6``
    (font + colour), styles tables (borders, header shading/bold, optional
    banding), shrinks the footnote/source list, and adds the page header/footer
    (with the optional ``logo`` — a :class:`LogoRender` — in the header).
    Tolerant of missing styles and malformed values (falls back to defaults).
    """
    from docx.shared import Pt, RGBColor

    body_font = styles.get("body_font") or STYLE_DEFAULTS["body_font"]
    heading_font = styles.get("heading_font") or STYLE_DEFAULTS["heading_font"]
    body_color = _norm_hex(styles.get("body_color"))
    heading_color = _norm_hex(styles.get("heading_color"))
    try:
        body_size = int(styles.get("body_size"))
    except (TypeError, ValueError):
        body_size = STYLE_DEFAULTS["body_size"]

    try:
        normal = doc.styles["Normal"]
        _set_style_font(normal, body_font)
        normal.font.size = Pt(body_size)
        if body_color:
            normal.font.color.rgb = RGBColor.from_string(body_color)
    except KeyError:
        pass

    for level in range(1, 7):
        try:
            heading = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        _set_style_font(heading, heading_font)
        if heading_color:
            heading.font.color.rgb = RGBColor.from_string(heading_color)

    _style_tables(doc, styles)
    _resize_footnote_sources(doc, max(1, body_size - FOOTNOTE_SIZE_DELTA))
    _apply_header_footer(doc, styles, logo)


# --- PDF export (CSS sibling of apply_doc_styles) ---------------------------

def _css_hex(value, default: str) -> str:
    """Return a ``#RRGGBB`` string from a stored hex value, or ``default``."""
    norm = _norm_hex(value)
    return f"#{norm}" if norm else default


def _css_string(text: str) -> str:
    """Escape a string for use inside a CSS ``content: "…"`` declaration."""
    return (text or "").replace("\\", "\\\\").replace('"', '\\"')


def _hf_box_css(styles: dict, prefix: str, content: str, family_stack: str) -> str:
    """Body of a ``@page`` margin box for the header/footer."""
    try:
        size = int(styles.get(f"{prefix}_size"))
    except (TypeError, ValueError):
        size = STYLE_DEFAULTS[f"{prefix}_size"]
    color = _css_hex(styles.get(f"{prefix}_color"), "#1A1A1A")
    weight = "700" if styles.get(f"{prefix}_bold") else "400"
    style = "italic" if styles.get(f"{prefix}_italic") else "normal"
    return (
        f"content:{content};"
        f"font-family:{family_stack};font-size:{size}pt;color:{color};"
        f"font-weight:{weight};font-style:{style};"
    )


def _logo_box_css(logo) -> str:
    """``@page`` margin-box body that paints the header logo at a fixed size.

    Embeds the bytes as a ``data:`` URL background (no temp file, like the fonts)
    and sets an explicit ``background-size`` computed from the aspect ratio — so
    the logo neither distorts nor exceeds the print box, independent of however
    WeasyPrint sizes the margin box itself. Painted to its side and vertically
    centered within the top margin.
    """
    width_cm, height_cm = _logo_print_dims(logo.aspect)
    b64 = base64.b64encode(logo.data).decode("ascii")
    side = "left" if logo.position == "left" else "right"
    return (
        'content:"";'
        f"background-image:url(data:image/png;base64,{b64});"
        "background-repeat:no-repeat;"
        f"background-position:{side} center;"
        f"background-size:{width_cm:.2f}cm {height_cm:.2f}cm;"
    )


def build_pdf_css(styles: dict, resolutions: dict, logo=None) -> str:
    """Build the stylesheet for PDF export — the CSS sibling of
    :func:`apply_doc_styles`, driven by the same ``styles`` dict so DOCX and PDF
    stay aligned.

    ``resolutions`` maps each font *name* in ``styles`` to a
    ``core.fonts.FontResolution`` (from ``resolve_fonts``). Fonts are embedded as
    ``@font-face`` data URLs; every font-family stack ends in an embedded face
    because the render host has Pango but no system font files. ``logo`` is an
    optional :class:`LogoRender` (from ``resolve_org_logo``) painted into a header
    margin box.
    """
    from core import fonts  # lazy: avoids a core.styles <-> core.fonts import cycle

    # @font-face blocks, deduped by css_family. Always include the generic
    # fallback face for each resolved font plus a mono face for code blocks.
    faces: dict[str, str] = {}

    def add(res) -> None:
        faces.setdefault(res.css_family, res.font_face_css())

    def stack(font_name: str) -> str:
        res = resolutions.get((font_name or "").strip())
        if res is None:
            res = fonts.bundled_resolution(FALLBACK_FONT)
        add(res)
        fb_key = fonts.GENERIC_FALLBACK.get(res.generic, FALLBACK_FONT)
        add(fonts.bundled_resolution(fb_key))
        return f"'{res.css_family}', 'wf-{fb_key.lower()}', {res.generic}"

    body_stack = stack(styles.get("body_font"))
    heading_stack = stack(styles.get("heading_font"))
    header_stack = stack(styles.get("header_font"))
    footer_stack = stack(styles.get("footer_font"))
    add(fonts.bundled_resolution("Cousine"))  # code blocks

    try:
        body_size = int(styles.get("body_size"))
    except (TypeError, ValueError):
        body_size = STYLE_DEFAULTS["body_size"]
    table_size = max(1, body_size - TABLE_SIZE_DELTA)
    body_color = _css_hex(styles.get("body_color"), "#1A1A1A")
    heading_color = _css_hex(styles.get("heading_color"), "#1A1A1A")
    accent = _norm_hex(styles.get("accent_color"))
    border_color = _css_hex(styles.get("table_border_color"), "#CCCCCC")
    border_style = styles.get("table_border_style") or "all"
    banded = bool(styles.get("table_banded"))
    band = _tint(accent, 0.85) if accent else "F2F2F2"

    # @page header/footer margin boxes (only when there's content to show).
    header_text = (styles.get("header_text") or "").strip()
    footer_text = (styles.get("footer_text") or "").strip()
    page_numbers = bool(styles.get("footer_page_numbers"))
    margin_boxes = ""
    # Header logo (if any) takes its side; header text takes the other. With no
    # logo, text stays in @top-left exactly as before.
    header_text_box = "@top-left"
    if logo is not None:
        logo_box = "@top-left" if logo.position == "left" else "@top-right"
        header_text_box = "@top-right" if logo.position == "left" else "@top-left"
        margin_boxes += logo_box + "{" + _logo_box_css(logo) + "}"
    if header_text:
        margin_boxes += header_text_box + "{" + _hf_box_css(
            styles, "header", f'"{_css_string(header_text)}"', header_stack
        ) + "}"
    if footer_text:
        margin_boxes += "@bottom-left{" + _hf_box_css(
            styles, "footer", f'"{_css_string(footer_text)}"', footer_stack
        ) + "}"
    if page_numbers:
        margin_boxes += "@bottom-right{" + _hf_box_css(
            styles, "footer", 'counter(page) " / " counter(pages)', footer_stack
        ) + "}"

    # Table borders per mode.
    if border_style == "all":
        table_border = f"table.wf,table.wf th,table.wf td{{border:0.5pt solid {border_color};}}"
    elif border_style == "horizontal":
        table_border = f"table.wf th,table.wf td{{border-top:0.5pt solid {border_color};border-bottom:0.5pt solid {border_color};}}"
    elif border_style == "header":
        table_border = f"table.wf thead th{{border-bottom:1pt solid {border_color};}}"
    else:
        table_border = ""

    accent_css = ""
    if accent:
        accent_css = f"table.wf thead th{{background:#{accent};color:#{_readable_text(accent)};}}"
    band_css = ""
    if banded:
        band_css = f"table.wf tbody tr:nth-child(even){{background:#{band};}}"

    return f"""
{"".join(faces.values())}
@page {{ size: A4; margin: 2.54cm; {margin_boxes} }}
html {{ font-family: {body_stack}; font-size: {body_size}pt; color: {body_color}; line-height: 1.45; }}
body {{ margin: 0; }}
p {{ margin: 0 0 0.6em; }}
h1, h2, h3, h4, h5, h6 {{ font-family: {heading_stack}; color: {heading_color}; line-height: 1.25; margin: 0.8em 0 0.4em; }}
h1 {{ font-size: 1.8em; }} h2 {{ font-size: 1.5em; }} h3 {{ font-size: 1.25em; }}
h4 {{ font-size: 1.1em; }} h5 {{ font-size: 1em; }} h6 {{ font-size: 0.9em; }}
ul, ol {{ margin: 0.4em 0 0.6em 1.4em; padding: 0; }}
li {{ margin: 0.15em 0; }}
blockquote {{ margin: 0.6em 0; padding-left: 10pt; border-left: 2pt solid #DDDDDD; color: #555555; }}
mark {{ background: #FFF3A3; }}
sup {{ font-size: 0.75em; vertical-align: super; line-height: 0; }}
code {{ font-family: 'wf-cousine', monospace; font-size: 0.9em; }}
pre {{ font-family: 'wf-cousine', monospace; font-size: 0.85em; background: #F5F5F5; padding: 6pt; border-radius: 3pt; white-space: pre-wrap; word-wrap: break-word; }}
pre code {{ background: none; font-size: inherit; }}
hr {{ border: none; border-bottom: 0.75pt solid #CCCCCC; margin: 10pt 0; height: 0; }}
img {{ max-width: 100%; height: auto; }}
table.wf {{ width: 100%; border-collapse: collapse; margin: 0.6em 0; font-size: {table_size}pt; }}
table.wf th, table.wf td {{ padding: 2pt 5pt; vertical-align: top; text-align: left; }}
table.wf thead {{ display: table-header-group; }}
table.wf th {{ font-weight: bold; }}
{table_border}
{accent_css}
{band_css}
""".strip()
