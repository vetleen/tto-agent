"""Tests for the org header-logo export path (core.styles)."""
import io
from types import SimpleNamespace

from django.test import SimpleTestCase
from docx import Document

from core import fonts
from core.styles import (
    LOGO_PRINT_MAX_H_CM,
    LOGO_PRINT_MAX_W_CM,
    STYLE_DEFAULTS,
    LogoRender,
    _logo_print_dims,
    apply_doc_styles,
    build_pdf_css,
    resolve_org_logo,
    validate_styles,
)


def _png_bytes(w, h):
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGBA", (w, h), (200, 30, 30, 255)).save(buf, format="PNG")
    return buf.getvalue()


class _FakeFieldFile:
    """Minimal stand-in for a FieldFile: truthy and openable as bytes."""

    def __init__(self, data):
        self._data = data

    def __bool__(self):
        return True

    def open(self, mode="rb"):
        return io.BytesIO(self._data)


class LogoPrintDimsTests(SimpleTestCase):
    def test_square_capped_by_height(self):
        w, h = _logo_print_dims(1.0)
        self.assertAlmostEqual(h, LOGO_PRINT_MAX_H_CM)
        self.assertAlmostEqual(w, LOGO_PRINT_MAX_H_CM)  # aspect 1 → w == h

    def test_moderate_wide_keeps_full_height(self):
        w, h = _logo_print_dims(2.0)
        self.assertAlmostEqual(h, LOGO_PRINT_MAX_H_CM)
        self.assertAlmostEqual(w, LOGO_PRINT_MAX_H_CM * 2)

    def test_very_wide_capped_by_width(self):
        w, h = _logo_print_dims(10.0)
        self.assertAlmostEqual(w, LOGO_PRINT_MAX_W_CM)
        self.assertAlmostEqual(h, LOGO_PRINT_MAX_W_CM / 10.0)
        self.assertLessEqual(h, LOGO_PRINT_MAX_H_CM)

    def test_degenerate_aspect_falls_back(self):
        w, h = _logo_print_dims(0)
        self.assertEqual((w, h), (LOGO_PRINT_MAX_W_CM, LOGO_PRINT_MAX_H_CM))


class ValidateLogoPositionTests(SimpleTestCase):
    def _valid(self, **over):
        data = dict(STYLE_DEFAULTS)
        data.update(over)
        return data

    def test_default_is_none(self):
        clean, err = validate_styles(self._valid())
        self.assertIsNone(err)
        self.assertEqual(clean["header_logo_position"], "none")

    def test_left_and_right_accepted(self):
        for pos in ("left", "right", "none"):
            clean, err = validate_styles(self._valid(header_logo_position=pos))
            self.assertIsNone(err)
            self.assertEqual(clean["header_logo_position"], pos)

    def test_bad_position_rejected(self):
        clean, err = validate_styles(self._valid(header_logo_position="middle"))
        self.assertIsNone(clean)
        self.assertIsNotNone(err)


class ResolveOrgLogoTests(SimpleTestCase):
    def test_none_position_returns_none(self):
        org = SimpleNamespace(logo=_FakeFieldFile(_png_bytes(100, 50)))
        styles = {**STYLE_DEFAULTS, "header_logo_position": "none"}
        self.assertIsNone(resolve_org_logo(styles, org))

    def test_no_logo_returns_none(self):
        org = SimpleNamespace(logo=None)
        styles = {**STYLE_DEFAULTS, "header_logo_position": "left"}
        self.assertIsNone(resolve_org_logo(styles, org))

    def test_resolves_render_with_aspect(self):
        org = SimpleNamespace(logo=_FakeFieldFile(_png_bytes(200, 100)))
        styles = {**STYLE_DEFAULTS, "header_logo_position": "left"}
        logo = resolve_org_logo(styles, org)
        self.assertIsNotNone(logo)
        self.assertEqual(logo.position, "left")
        self.assertAlmostEqual(logo.aspect, 2.0, places=3)

    def test_unreadable_logo_returns_none(self):
        org = SimpleNamespace(logo=_FakeFieldFile(b"not an image"))
        styles = {**STYLE_DEFAULTS, "header_logo_position": "right"}
        self.assertIsNone(resolve_org_logo(styles, org))


class BuildPdfCssLogoTests(SimpleTestCase):
    def _css(self, logo, **over):
        styles = {**STYLE_DEFAULTS, "header_text": "Hi", **over}
        return build_pdf_css(styles, fonts.resolve_fonts(styles, None), logo)

    def test_no_logo_keeps_text_top_left(self):
        css = build_pdf_css(
            {**STYLE_DEFAULTS, "header_text": "Hi"},
            fonts.resolve_fonts(STYLE_DEFAULTS, None),
        )
        self.assertIn('@top-left{content:"Hi"', css)
        self.assertNotIn("background-image", css)

    def test_logo_left_text_right(self):
        logo = LogoRender(data=_png_bytes(200, 100), aspect=2.0, position="left")
        css = self._css(logo, header_logo_position="left")
        self.assertIn('@top-left{content:"";background-image:url(data:image/png;base64,', css)
        self.assertIn("background-size:", css)
        self.assertIn('@top-right{content:"Hi"', css)

    def test_logo_box_has_explicit_width(self):
        """The logo margin box needs an explicit width or it collapses to zero
        (empty content) and WeasyPrint paints no background — the bug that left
        the logo out of the PDF while it showed in the DOCX. aspect 2.0 →
        height 1.4cm, width 2.8cm."""
        logo = LogoRender(data=_png_bytes(200, 100), aspect=2.0, position="left")
        css = self._css(logo, header_logo_position="left")
        self.assertIn("width:2.80cm;", css)
        self.assertIn("background-size:2.80cm 1.40cm;", css)

    def test_logo_right_text_left(self):
        logo = LogoRender(data=_png_bytes(200, 100), aspect=2.0, position="right")
        css = self._css(logo, header_logo_position="right")
        self.assertIn('@top-right{content:"";background-image:url(data:image/png;base64,', css)
        self.assertIn('@top-left{content:"Hi"', css)


class ApplyDocStylesLogoTests(SimpleTestCase):
    def _header_p(self, doc):
        return doc.sections[0].header.paragraphs[0]

    def test_logo_left_with_text_adds_picture_and_tab(self):
        doc = Document()
        doc.add_paragraph("Body")
        logo = LogoRender(data=_png_bytes(200, 100), aspect=2.0, position="left")
        apply_doc_styles(doc, {**STYLE_DEFAULTS, "header_text": "Acme"}, logo)
        hp = self._header_p(doc)
        self.assertEqual(len(hp._p.xpath(".//w:drawing")), 1)  # picture embedded
        self.assertEqual(len(hp._p.xpath(".//w:tab")), 1)      # text pushed to other side
        self.assertIn("Acme", hp.text)

    def test_logo_only_no_text(self):
        doc = Document()
        doc.add_paragraph("Body")
        logo = LogoRender(data=_png_bytes(300, 100), aspect=3.0, position="left")
        apply_doc_styles(doc, {**STYLE_DEFAULTS, "header_text": ""}, logo)
        hp = self._header_p(doc)
        self.assertEqual(len(hp._p.xpath(".//w:drawing")), 1)
        self.assertEqual(hp.text.strip(), "")

    def test_no_logo_no_picture(self):
        doc = Document()
        doc.add_paragraph("Body")
        apply_doc_styles(doc, {**STYLE_DEFAULTS, "header_text": "Acme"})
        hp = self._header_p(doc)
        self.assertEqual(len(hp._p.xpath(".//w:drawing")), 0)
        self.assertEqual(hp.text, "Acme")
