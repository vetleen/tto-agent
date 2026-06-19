"""Tests for core.styles — org document-export design levers."""
from types import SimpleNamespace

from django.test import SimpleTestCase
from docx import Document
from docx.oxml.ns import qn

from core.styles import (
    FOOTNOTE_MARKER,
    STYLE_DEFAULTS,
    apply_doc_styles,
    get_org_styles,
    validate_styles,
)


class GetOrgStylesTests(SimpleTestCase):
    def test_none_org_returns_defaults_copy(self):
        result = get_org_styles(None)
        self.assertEqual(result, STYLE_DEFAULTS)
        self.assertIsNot(result, STYLE_DEFAULTS)  # never hand back the module dict

    def test_partial_styles_merged_over_defaults(self):
        org = SimpleNamespace(preferences={"styles": {"body_font": "Georgia", "body_size": 13}})
        styles = get_org_styles(org)
        self.assertEqual(styles["body_font"], "Georgia")
        self.assertEqual(styles["body_size"], 13)
        self.assertEqual(styles["heading_font"], STYLE_DEFAULTS["heading_font"])

    def test_no_styles_key_returns_defaults(self):
        org = SimpleNamespace(preferences={})
        self.assertEqual(get_org_styles(org), STYLE_DEFAULTS)


class ValidateStylesTests(SimpleTestCase):
    def _valid(self, **over):
        data = {
            "body_font": "Georgia",
            "body_size": 12,
            "heading_font": "Cambria",
            "heading_color": "#112233",
            "body_color": "#000000",
            "accent_color": "#2563EB",
            "table_border_style": "horizontal",
            "table_border_color": "#999999",
            "table_banded": True,
        }
        data.update(over)
        return data

    def test_valid_payload(self):
        clean, err = validate_styles(self._valid())
        self.assertIsNone(err)
        self.assertEqual(clean["body_font"], "Georgia")
        self.assertEqual(clean["heading_color"], "#112233")
        self.assertEqual(clean["body_size"], 12)

    def test_custom_font_accepted(self):
        clean, err = validate_styles(self._valid(body_font="PT Sans"))
        self.assertIsNone(err)
        self.assertEqual(clean["body_font"], "PT Sans")

    def test_blank_accent_allowed(self):
        clean, err = validate_styles(self._valid(accent_color=""))
        self.assertIsNone(err)
        self.assertEqual(clean["accent_color"], "")

    def test_lowercase_hex_normalised(self):
        clean, err = validate_styles(self._valid(accent_color="#abcdef"))
        self.assertIsNone(err)
        self.assertEqual(clean["accent_color"], "#ABCDEF")

    def test_bad_hex_rejected(self):
        clean, err = validate_styles(self._valid(heading_color="blue"))
        self.assertIsNone(clean)
        self.assertIsNotNone(err)

    def test_size_out_of_range_rejected(self):
        self.assertIsNone(validate_styles(self._valid(body_size=40))[0])
        self.assertIsNone(validate_styles(self._valid(body_size=2))[0])

    def test_illegal_font_char_rejected(self):
        self.assertIsNone(validate_styles(self._valid(body_font="Bad<Font>"))[0])

    def test_table_levers_round_trip(self):
        clean, err = validate_styles(self._valid())
        self.assertIsNone(err)
        self.assertEqual(clean["table_border_style"], "horizontal")
        self.assertEqual(clean["table_border_color"], "#999999")
        self.assertIs(clean["table_banded"], True)

    def test_bad_border_style_rejected(self):
        self.assertIsNone(validate_styles(self._valid(table_border_style="dotted"))[0])

    def test_bad_border_color_rejected(self):
        self.assertIsNone(validate_styles(self._valid(table_border_color="grey"))[0])

    def test_banded_must_be_bool(self):
        self.assertIsNone(validate_styles(self._valid(table_banded="yes"))[0])

    def test_missing_keys_fall_back_to_defaults(self):
        clean, err = validate_styles({})
        self.assertIsNone(err)
        self.assertEqual(clean, STYLE_DEFAULTS)

    def test_non_dict_rejected(self):
        self.assertIsNone(validate_styles("nope")[0])


class ApplyDocStylesTests(SimpleTestCase):
    def test_body_and_heading_styles_applied(self):
        doc = Document()
        doc.add_paragraph("body text")
        doc.add_heading("A heading", level=1)
        apply_doc_styles(doc, {
            **STYLE_DEFAULTS,
            "body_font": "Georgia",
            "body_size": 13,
            "heading_font": "Cambria",
            "heading_color": "#112233",
            "body_color": "#445566",
        })
        normal = doc.styles["Normal"]
        self.assertEqual(normal.font.name, "Georgia")
        self.assertEqual(normal.font.size.pt, 13)
        self.assertEqual(str(normal.font.color.rgb), "445566")
        h1 = doc.styles["Heading 1"]
        self.assertEqual(h1.font.name, "Cambria")
        self.assertEqual(str(h1.font.color.rgb), "112233")
        # Heading styles reference a theme font (w:asciiTheme) that Word prefers
        # over an explicit w:ascii — assert it's stripped so the font renders.
        rfonts = h1.element.get_or_add_rPr().get_or_add_rFonts()
        self.assertEqual(rfonts.get(qn("w:ascii")), "Cambria")
        self.assertIsNone(rfonts.get(qn("w:asciiTheme")))

    def test_accent_shades_first_table_row(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header"
        apply_doc_styles(doc, {**STYLE_DEFAULTS, "accent_color": "#222222"})
        tcPr = table.rows[0].cells[0]._tc.find(qn("w:tcPr"))
        self.assertIsNotNone(tcPr)
        shd = tcPr.find(qn("w:shd"))
        self.assertIsNotNone(shd)
        self.assertEqual(shd.get(qn("w:fill")), "222222")

    def test_blank_accent_leaves_tables_unshaded(self):
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        apply_doc_styles(doc, {**STYLE_DEFAULTS, "accent_color": ""})
        tcPr = table.rows[0].cells[0]._tc.find(qn("w:tcPr"))
        shd = tcPr.find(qn("w:shd")) if tcPr is not None else None
        self.assertIsNone(shd)

    def test_malformed_size_falls_back_to_default(self):
        doc = Document()
        apply_doc_styles(doc, {**STYLE_DEFAULTS, "body_size": "huge"})
        self.assertEqual(doc.styles["Normal"].font.size.pt, STYLE_DEFAULTS["body_size"])

    @staticmethod
    def _cell_shd(cell):
        tcPr = cell._tc.find(qn("w:tcPr"))
        return None if tcPr is None else tcPr.find(qn("w:shd"))

    def test_table_borders_applied(self):
        doc = Document()
        doc.add_table(rows=2, cols=2)
        apply_doc_styles(doc, {**STYLE_DEFAULTS, "table_border_style": "all", "table_border_color": "#999999"})
        borders = doc.tables[0]._tbl.tblPr.find(qn("w:tblBorders"))
        self.assertIsNotNone(borders)
        self.assertIsNotNone(borders.find(qn("w:insideH")))
        self.assertEqual(borders.find(qn("w:top")).get(qn("w:color")), "999999")

    def test_table_border_style_none(self):
        doc = Document()
        doc.add_table(rows=2, cols=2)
        apply_doc_styles(doc, {**STYLE_DEFAULTS, "table_border_style": "none"})
        self.assertIsNone(doc.tables[0]._tbl.tblPr.find(qn("w:tblBorders")))

    def test_header_bold_and_repeats(self):
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Head"
        apply_doc_styles(doc, {**STYLE_DEFAULTS})
        self.assertTrue(table.rows[0].cells[0].paragraphs[0].runs[0].font.bold)
        trPr = table.rows[0]._tr.find(qn("w:trPr"))
        self.assertIsNotNone(trPr.find(qn("w:tblHeader")))

    def test_footnote_sources_shrunk_two_points_below_body(self):
        doc = Document()
        marked = doc.add_paragraph(FOOTNOTE_MARKER + "BCC Research", style="List Number")
        plain = doc.add_paragraph("Regular numbered item", style="List Number")
        apply_doc_styles(doc, {**STYLE_DEFAULTS, "body_size": 11})
        # marked source item: shrunk to 9pt and marker stripped
        self.assertEqual(marked.runs[0].font.size.pt, 9)
        self.assertNotIn(FOOTNOTE_MARKER, marked.text)
        # a normal numbered list is left at the inherited (unset) size
        self.assertIsNone(plain.runs[0].font.size)

    def test_banded_rows_shade_alternate_body_rows(self):
        doc = Document()
        table = doc.add_table(rows=4, cols=1)
        apply_doc_styles(doc, {**STYLE_DEFAULTS, "table_banded": True})
        self.assertIsNone(self._cell_shd(table.rows[1].cells[0]))   # 1st body row
        self.assertIsNotNone(self._cell_shd(table.rows[2].cells[0]))  # 2nd body row banded
        self.assertIsNone(self._cell_shd(table.rows[3].cells[0]))   # 3rd body row
