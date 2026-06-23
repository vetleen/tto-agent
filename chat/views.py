import io
import json
import logging
import re

from django.contrib.auth.decorators import login_required
from django.http import FileResponse, Http404, HttpResponseBadRequest, JsonResponse
from django.shortcuts import get_object_or_404, render
from django.views.decorators.http import require_http_methods, require_POST

from chat.models import ChatAttachment, ChatCanvas, ChatThread, ThreadTask
from documents.models import DataRoom, DataRoomDocument, DataRoomDocumentTag

logger = logging.getLogger(__name__)


def _get_accessible_data_rooms(user):
    """Return data rooms the user owns, each annotated with its active
    (non-archived) document count for display in the attach modal."""
    from django.db.models import Count, Q

    return list(
        DataRoom.objects.filter(created_by=user, is_archived=False)
        .annotate(
            document_count=Count("documents", filter=Q(documents__is_archived=False))
        )
        .order_by("-updated_at")
        .values("pk", "uuid", "name", "document_count")
    )


def _user_can_access_image_asset(user, asset) -> bool:
    """Re-derive access from the asset's owner (version / canvas / message / thread).

    Returns False for orphans. This is the only gate on serving image bytes —
    never a presigned S3 URL.
    """
    if asset.canvas_id:
        return asset.canvas.thread.created_by_id == user.id
    if asset.message_id:
        return asset.message.thread.created_by_id == user.id
    if asset.thread_id:
        return asset.thread.created_by_id == user.id
    if asset.version_id:
        from documents.views import _user_can_access_data_room

        return _user_can_access_data_room(user, asset.version.document.data_room)
    return False


@login_required
@require_http_methods(["GET"])
def serve_image_asset(request, asset_id):
    """Stream an ImageAsset's bytes when the user may access its owner.

    Access is re-checked on every request. Only known-safe image types are
    served inline (so an <img> can render them); anything else is forced to
    download with a generic content type.
    """
    from chat.image_assets import image_asset_source
    from chat.models import ImageAsset
    from core.file_types import KIND_IMAGE, kind_for_mime

    asset = get_object_or_404(ImageAsset, id=asset_id)
    if not _user_can_access_image_asset(request.user, asset):
        raise Http404

    # A reference asset (empty blob) resolves to the data-room version's native
    # image; a normal asset serves its own blob.
    source, ct = image_asset_source(asset)
    if source is None:
        raise Http404
    displayable = kind_for_mime(ct) == KIND_IMAGE
    resp = FileResponse(
        source.open("rb"),
        content_type=ct if displayable else "application/octet-stream",
    )
    disposition = "inline" if displayable else "attachment"
    resp["Content-Disposition"] = f'{disposition}; filename="{asset.id}"'
    resp["X-Content-Type-Options"] = "nosniff"
    return resp


_IMAGE_TOKEN_RE = re.compile(r"\[\[image:([0-9a-fA-F-]{36})\|([^\]]*)\]\]")


def _embed_image_tokens(content: str, user) -> str:
    """Replace ``[[image:<uuid>|label]]`` tokens with ``<img>`` data-URLs for
    docx export. Tokens the user can't access (or that no longer resolve) become
    a neutral placeholder instead of vanishing. Runs server-side (sync)."""
    import base64

    from chat.image_assets import IMAGE_UNAVAILABLE_TEXT, image_asset_source
    from chat.models import ImageAsset

    placeholder = f"<em>{IMAGE_UNAVAILABLE_TEXT}</em>"

    def repl(m):
        uuid, label = m.group(1), m.group(2)
        try:
            asset = ImageAsset.objects.get(id=uuid)
        except (ImageAsset.DoesNotExist, ValueError):
            return placeholder
        if not _user_can_access_image_asset(user, asset):
            return placeholder
        source, ct = image_asset_source(asset)
        if source is None:
            return placeholder
        try:
            with source.open("rb") as f:
                data = f.read()
        except Exception:
            return placeholder
        b64 = base64.b64encode(data).decode("ascii")
        alt = (label or "").replace('"', "'").replace("<", "").replace(">", "")
        return f'<img src="data:{ct or "image/png"};base64,{b64}" alt="{alt}" />'

    return _IMAGE_TOKEN_RE.sub(repl, content)


@login_required
@require_POST
def reattach_attachment(request, thread_id, attachment_id):
    """Re-attach a previously-uploaded file to the user's next message.

    Copies the stored bytes into a fresh ChatAttachment (message=NULL) so it
    rides along on the next message like a new upload. We copy rather than share
    the file because the per-attachment file-delete signal would otherwise take
    the original's bytes with it.
    """
    from django.core.files.base import ContentFile

    thread = get_object_or_404(ChatThread, id=thread_id, created_by=request.user)
    old = get_object_or_404(ChatAttachment, id=attachment_id, thread=thread)
    try:
        with old.file.open("rb") as f:
            data = f.read()
    except Exception:
        return JsonResponse({"error": "That file is no longer available."}, status=404)

    new = ChatAttachment(
        thread=thread,
        uploaded_by=request.user,
        original_filename=old.original_filename,
        content_type=old.content_type,
        size_bytes=old.size_bytes,
    )
    new.file.save(old.original_filename[:255] or "file", ContentFile(data), save=True)
    return JsonResponse({
        "id": str(new.id),
        "filename": new.original_filename,
        "content_type": new.content_type,
        "size_bytes": new.size_bytes,
    })


def _group_threads_by_time(threads):
    """Bucket threads (already ordered by ``-updated_at``) into Today / Yesterday /
    Earlier by their ``updated_at`` in the active local timezone.

    Returns a list of ``{"label", "threads"}`` dicts, omitting empty groups. The
    sidebar always presents the Today group first; JS recreates it when a new or
    just-touched thread needs to land there.
    """
    from datetime import timedelta

    from django.utils import timezone

    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    today_items, yesterday_items, earlier_items = [], [], []
    for t in threads:
        dt = t.updated_at
        if timezone.is_aware(dt):
            dt = timezone.localtime(dt)
        d = dt.date()
        if d >= today:
            today_items.append(t)
        elif d == yesterday:
            yesterday_items.append(t)
        else:
            earlier_items.append(t)
    groups = []
    if today_items:
        groups.append({"label": "Today", "threads": today_items})
    if yesterday_items:
        groups.append({"label": "Yesterday", "threads": yesterday_items})
    if earlier_items:
        groups.append({"label": "Earlier", "threads": earlier_items})
    return groups


@login_required
@require_http_methods(["GET"])
def chat_home(request):
    """Main chat page. Loads a thread via ?thread=<uuid> if provided."""
    thread = None
    chat_messages = []
    thread_loop_id = None
    history_compressed_above = False

    if request.GET.get("thread"):
        thread_id = request.GET["thread"]
        thread = ChatThread.objects.filter(
            id=thread_id, created_by=request.user
        ).first()
        if thread:
            from django.utils import timezone

            from chat.models import Loop

            loop = Loop.objects.filter(thread=thread).first()
            # Auto-restore archived threads when opened — except paused loops,
            # whose archived state mirrors their paused state (revive them from
            # the Loops page or the restore action, not by merely opening).
            if thread.is_archived and not (
                loop and loop.status == Loop.Status.PAUSED
            ):
                thread.is_archived = False
                thread.save(update_fields=["is_archived"])
            # If this thread is a loop, mark it seen (clears the unread badge) and
            # expose its id so /loop in this thread opens the edit modal.
            if loop:
                loop.last_seen_at = timezone.now()
                loop.save(update_fields=["last_seen_at"])
                thread_loop_id = str(loop.id)
            # Visible messages, plus hidden assistant tool-loop messages that
            # carry narration or thinking — those render as collapsed
            # "Thought further" blocks (matching the live-streaming UI).
            from django.db.models import Q

            chat_messages = list(
                thread.messages.filter(
                    Q(is_hidden_from_user=False)
                    | Q(is_hidden_from_user=True, role="assistant", is_redacted=False)
                ).order_by("created_at")[:100]
            )
            filtered = []
            for m in chat_messages:
                if m.is_hidden_from_user:
                    has_narration = bool(m.content.strip()) or bool(
                        (m.metadata or {}).get("thinking")
                    )
                    if not has_narration:
                        continue
                    m.is_intermediate = True
                filtered.append(m)
            chat_messages = filtered

            # Annotate user messages with attachments (id used for re-attach)
            msg_ids = [m.pk for m in chat_messages if m.role == "user"]
            if msg_ids:
                from collections import defaultdict

                from chat.models import ImageAsset
                att_qs = ChatAttachment.objects.filter(
                    message_id__in=msg_ids
                ).values("id", "message_id", "original_filename", "content_type")
                att_map = defaultdict(list)
                for a in att_qs:
                    att_map[a["message_id"]].append({
                        "id": str(a["id"]),
                        "name": a["original_filename"],
                        "content_type": a["content_type"],
                    })
                # Embedded images extracted from a user's docx/pdf attachments are
                # persisted as ImageAssets scoped to that user message; surface
                # them as viewable thumbnails (served via chat_image_asset).
                img_qs = ImageAsset.objects.filter(
                    message_id__in=msg_ids
                ).values("id", "message_id", "description")
                img_map = defaultdict(list)
                for a in img_qs:
                    img_map[a["message_id"]].append({
                        "id": str(a["id"]),
                        "description": a["description"],
                    })
                for m in chat_messages:
                    m.attachment_items = att_map.get(m.pk, [])
                    m.attachment_images = img_map.get(m.pk, [])

            # Mark which shown messages fall before the compression boundary, so
            # the UI can show a divider and flag attachments that are no longer
            # re-sent to the model on each turn.
            if thread.summary_up_to_message_id:
                from chat.models import ChatMessage

                boundary = ChatMessage.objects.filter(
                    pk=thread.summary_up_to_message_id
                ).only("created_at").first()
                boundary_dt = boundary.created_at if boundary else None
                if boundary_dt:
                    summarized = []
                    for m in chat_messages:
                        m.is_summarized = m.created_at <= boundary_dt
                        if m.is_summarized:
                            summarized.append(m)
                    if summarized:
                        summarized[-1].is_boundary = True
                    elif chat_messages:
                        history_compressed_above = True

    all_threads = ChatThread.objects.filter(created_by=request.user).order_by("-updated_at")
    threads = [t for t in all_threads if not t.is_archived]
    archived_threads = [t for t in all_threads if t.is_archived]

    # Annotate threads that are loops so the sidebar can show a loop icon + an
    # unread dot when a scheduled run produced a result the user hasn't opened.
    # Paused loops live in the Archived section, so annotate those too.
    from chat.models import Loop

    loop_map = {
        loop.thread_id: loop
        for loop in Loop.objects.filter(thread__in=all_threads)
    }
    for t in [*threads, *archived_threads]:
        loop = loop_map.get(t.id)
        t.is_loop = loop is not None
        t.loop_unread = loop.is_unread if loop else False
        t.loop_id = str(loop.id) if loop else ""
        t.loop_paused = (loop.status == Loop.Status.PAUSED) if loop else False

    thread_groups = _group_threads_by_time(threads)

    # User's data rooms (owned + shared) for the attach modal
    data_rooms = _get_accessible_data_rooms(request.user)

    # If thread selected, get its tasks
    thread_tasks_json = "[]"
    if thread:
        thread_tasks = list(
            ThreadTask.objects.filter(thread=thread)
            .order_by("order", "created_at")
            .values("id", "title", "status")
        )
        if thread_tasks:
            # Convert UUIDs to strings for JSON serialization
            for t in thread_tasks:
                t["id"] = str(t["id"])
            thread_tasks_json = json.dumps(thread_tasks)

    # Compute thread cost
    thread_cost_usd = 0.0
    if thread:
        from django.db.models import Sum

        from llm.models import LLMCallLog

        result = LLMCallLog.objects.filter(
            conversation_id=str(thread.id),
        ).aggregate(total=Sum("cost_usd"))
        thread_cost_usd = float(result["total"]) if result["total"] is not None else 0.0

    # Active sub-agents for status bar
    active_subagent_count = 0
    if thread:
        from chat.models import SubAgentRun

        active_subagent_count = SubAgentRun.objects.filter(
            thread_id=thread.id,
            status__in=[SubAgentRun.Status.PENDING, SubAgentRun.Status.RUNNING],
        ).count()

    # If thread selected, get its attached data rooms
    thread_data_rooms = []
    thread_skills = []
    if thread:
        from chat.models import ChatThreadSkill

        thread_data_rooms = list(
            thread.data_rooms.values("pk", "uuid", "name")
        )
        thread_skills = [
            {
                "id": str(r.skill_id),
                "name": r.skill.name,
                "emoji": r.skill.emoji,
            }
            for r in ChatThreadSkill.objects.filter(
                thread=thread, skill__is_active=True
            ).select_related("skill")
        ]

    # Resolve preferences (model choices, allowed skills, etc.)
    from django.conf import settings as django_settings

    from core.file_types import CHAT_KINDS, accept_attr
    from core.preferences import get_preferences
    from llm.display import (
        get_capability_level,
        get_display_name,
        get_model_meta_tooltip,
        get_price_level,
        get_thinking_levels,
        supports_thinking,
        supports_vision,
    )

    prefs = get_preferences(request.user)

    # Skills for the skill modal — use prefs.allowed_skills which respects
    # org admin enable/disable settings, rather than raw get_available_skills()
    skills_json = json.dumps([
        {
            "id": s["id"],
            "name": s["name"],
            "emoji": s.get("emoji", ""),
            "description": s.get("description", ""),
        }
        for s in prefs.allowed_skills
    ])
    model_choices = [
        {
            "id": m,
            "display_name": get_display_name(m),
            "supports_thinking": supports_thinking(m),
            "supports_vision": supports_vision(m),
            "thinking_levels": get_thinking_levels(m),
            "price_level": get_price_level(m),
            "capability_level": get_capability_level(m),
            "meta_tooltip": get_model_meta_tooltip(m),
        }
        for m in prefs.allowed_models
    ]

    pending_initial_turn = bool(
        thread and (thread.metadata or {}).get("pending_initial_turn")
    )

    # Model picker default: the loaded thread's effective model (honoring its
    # stored choice, with tier fallback), or the user's preferred chat model for
    # a new thread.
    from core.preferences import resolve_thread_model

    preferred_chat_model = prefs.feature_models.get("chat", prefs.top_model)
    effective_model = (
        resolve_thread_model(thread.model, prefs) if thread else preferred_chat_model
    )

    return render(
        request,
        "chat/chat.html",
        {
            "thread": thread,
            "thread_loop_id": thread_loop_id,
            "threads": threads,
            "thread_groups": thread_groups,
            "archived_threads": archived_threads,
            "messages": chat_messages,
            "history_compressed_above": history_compressed_above,
            "data_rooms": data_rooms,
            "thread_data_rooms": thread_data_rooms,
            "skills_json": skills_json,
            "thread_skills_json": json.dumps(thread_skills),
            "model_choices_json": json.dumps(model_choices),
            "default_model": effective_model,
            "default_model_display": get_display_name(effective_model),
            "preferred_chat_model": preferred_chat_model,
            "preferred_chat_model_display": get_display_name(preferred_chat_model),
            "thread_tasks_json": thread_tasks_json,
            "thread_cost_usd": thread_cost_usd,
            "pending_initial_turn": pending_initial_turn,
            "allow_agent_attach_skills": prefs.allow_agent_attach_skills,
            "active_subagent_count": active_subagent_count,
            "assistant_name": django_settings.ASSISTANT_NAME,
            # File-picker accept list derived from the unified capability table.
            # Includes image/* so iOS Safari offers the photo library instead of
            # defaulting to the camera/video flow.
            "attach_accept": accept_attr(CHAT_KINDS),
        },
    )


@login_required
@require_POST
def thread_delete(request, thread_id):
    thread = get_object_or_404(
        ChatThread, id=thread_id, created_by=request.user
    )
    thread.delete()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def thread_archive(request, thread_id):
    thread = get_object_or_404(
        ChatThread, id=thread_id, created_by=request.user
    )
    thread.is_archived = not thread.is_archived
    thread.save(update_fields=["is_archived"])
    # A loop thread's archived state mirrors its paused state: archiving a live
    # loop pauses it; restoring a paused loop restarts it.
    from chat.models import Loop

    loop = Loop.objects.filter(thread=thread).first()
    if loop is not None:
        from django.utils import timezone

        from chat.loop_service import pause_loop, restart_loop

        loop.thread = thread  # reuse the already-saved instance
        if thread.is_archived and loop.status == Loop.Status.ACTIVE:
            pause_loop(loop)
        elif not thread.is_archived and loop.status == Loop.Status.PAUSED:
            restart_loop(loop, timezone.now())
    return JsonResponse({"ok": True, "is_archived": thread.is_archived})


@login_required
@require_POST
def thread_rename(request, thread_id):
    thread = get_object_or_404(
        ChatThread, id=thread_id, created_by=request.user
    )
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)
    title = body.get("title", "").strip()[:255]
    if not title:
        return JsonResponse({"error": "Please enter a title."}, status=400)
    thread.title = title
    thread.save(update_fields=["title"])
    return JsonResponse({"ok": True, "title": thread.title})


@login_required
@require_POST
def thread_emoji(request, thread_id):
    thread = get_object_or_404(
        ChatThread, id=thread_id, created_by=request.user
    )
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)
    emoji = body.get("emoji", "")[:8]
    thread.emoji = emoji
    thread.save(update_fields=["emoji"])
    return JsonResponse({"ok": True, "emoji": thread.emoji})


def _fix_table_trailing_content(doc):
    """Repair html2docx's broken output after each table.

    html2docx emits several blank ``" "`` paragraphs after a table and leaks the
    next block's text into the last one; when that block is a heading it also
    leaves the heading element empty, so the heading silently renders as body
    text. For each table we: (1) move a heading's text back out of the Normal
    paragraph it leaked into, (2) drop the spurious blanks, and (3) keep exactly
    one empty paragraph after the table — for spacing, and so Word won't merge
    two otherwise-adjacent tables. Image/drawing paragraphs are preserved.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    body = doc.element.body
    P, TBL = qn("w:p"), qn("w:tbl")

    def is_blank(p):
        if p.findall(".//" + qn("w:drawing")) or p.findall(".//" + qn("w:pict")):
            return False
        return "".join(t.text or "" for t in p.iter(qn("w:t"))).strip() == ""

    def is_heading(p):
        pPr = p.find(qn("w:pPr"))
        ps = pPr.find(qn("w:pStyle")) if pPr is not None else None
        return ps is not None and (ps.get(qn("w:val")) or "").startswith("Heading")

    def lstrip_para(p):
        for t in p.iter(qn("w:t")):
            t.text = (t.text or "").lstrip()
            if t.text:
                break

    for table in list(doc.tables):
        tbl = table._tbl

        # Collect the run of paragraphs immediately following the table.
        run = []
        sib = tbl.getnext()
        while sib is not None and sib.tag == P:
            run.append(sib)
            sib = sib.getnext()

        # Split into leading blanks and the first text-bearing ("leaked") para.
        blanks, leaked = [], None
        for p in run:
            if is_blank(p):
                blanks.append(p)
            else:
                leaked = p
                break

        if leaked is not None:
            after = leaked.getnext()
            if (after is not None and after.tag == P and is_blank(after)
                    and is_heading(after) and not is_heading(leaked)):
                # Heading text leaked into a Normal paragraph — move it back.
                for r in leaked.findall(qn("w:r")):
                    after.append(r)
                lstrip_para(after)
                body.remove(leaked)
            else:
                lstrip_para(leaked)

        for blank in blanks:
            body.remove(blank)

        # One spacer after the table (skip if the table ends the document).
        if leaked is not None or (sib is not None and sib.tag == TBL):
            tbl.addnext(OxmlElement("w:p"))


# html2docx drops <hr>, so we swap each one for a paragraph carrying this
# private-use marker, then render it as a divider rule in _render_dividers.
_HR_DIVIDER_MARKER = "\uE000"  # private-use char, never in real content
_HR_RE = re.compile(r"<hr\s*/?>")


def _render_dividers(doc):
    """Turn the assistant's ``---`` thematic breaks into a divider rule.

    Markdown ``---`` becomes ``<hr>``, which html2docx silently drops. We
    replaced each one with a marker paragraph (see ``canvas_export``); here we
    clear the marker and render the paragraph as a thin light-grey rule with
    breathing room above and below — a clean section break.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    for para in doc.paragraphs:
        if para.text.strip() != _HR_DIVIDER_MARKER:
            continue
        for run in list(para.runs):
            run._r.getparent().remove(run._r)
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(10)
        pPr = para._p.get_or_add_pPr()
        old = pPr.find(qn("w:pBdr"))
        if old is not None:
            pPr.remove(old)
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # ~0.75pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        # pBdr precedes shd/spacing/ind/jc/rPr in the pPr schema order.
        pPr.insert_element_before(pBdr, "w:shd", "w:tabs", "w:spacing", "w:ind",
                                  "w:jc", "w:rPr", "w:sectPr")


def _strip_space_after_line_breaks(doc):
    """Drop the leading space html2docx adds after a markdown hard break.

    A markdown hard break becomes ``<br/>`` followed by the source newline;
    html2docx turns that newline into a leading space on the continued line, so
    a hard-wrapped block (e.g. a "Prepared for: / Platform: / Status:" header)
    renders with a stray space at the start of every line but the first. We
    strip the whitespace of the first text run after each break.
    """
    from docx.oxml.ns import qn

    BR, T = qn("w:br"), qn("w:t")
    for para in doc.element.body.iter(qn("w:p")):
        after_break = False
        for el in para.iter():
            if el.tag == BR:
                after_break = True
            elif el.tag == T:
                if after_break and el.text:
                    el.text = el.text.lstrip()
                after_break = False


@login_required
@require_http_methods(["POST"])
async def canvas_export(request, thread_id, canvas_id=None):
    """Export the canvas as a .docx file.

    The client sends JSON ``{"content": "..."}`` with mermaid blocks already
    rendered as ``<img>`` tags (rendered in the browser), so the server never
    runs a headless browser over untrusted canvas content.
    """
    from asgiref.sync import sync_to_async

    thread = await sync_to_async(get_object_or_404)(ChatThread, id=thread_id, created_by=request.user)
    if canvas_id:
        canvas = await sync_to_async(get_object_or_404)(ChatCanvas, pk=canvas_id, thread=thread)
    else:
        canvas = await sync_to_async(get_object_or_404)(ChatCanvas, pk=thread.active_canvas_id, thread=thread)

    import markdown as md

    from chat.markdown_export import MarkExtension, html2docx_with_highlight as html2docx
    from chat.services import normalize_heading_levels, render_citations, replace_email_with_html

    # Client already replaced mermaid blocks with <img> tags.
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return HttpResponseBadRequest("Invalid JSON")
    content = body.get("content", canvas.content)

    # Promote headings so the doc's top level is H1 (LLMs often start at ##).
    content = normalize_heading_levels(content)
    # Footnote citations [^1] → superscripts + numbered Sources list.
    content = render_citations(content)
    content = replace_email_with_html(content)
    # Resolve image-asset tokens to embedded <img> data-URLs (same path mermaid
    # images take), so exported docs keep their images.
    content = await sync_to_async(_embed_image_tokens)(content, request.user)
    html_content = md.markdown(content, extensions=["tables", "fenced_code", MarkExtension()])
    # html2docx drops <hr>; swap each for a marker paragraph we style as a rule.
    html_content = _HR_RE.sub(f"<p>{_HR_DIVIDER_MARKER}</p>", html_content)
    full_html = f"<html><body>{html_content}</body></html>"
    buf = html2docx(full_html, title=canvas.title)

    # html2docx sets gridCol widths to span the full page but leaves
    # cell widths at 0 and table width at auto. With autofit Word
    # auto-sizes from content, making tables narrower than the page.
    # Fix: set table width to 100%, copy gridCol widths into cells,
    # and switch to fixed layout so Word honours the widths exactly.
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    doc = DocxDocument(buf)
    for table in doc.tables:
        table.autofit = False
        tblW = table._tbl.tblPr.find(qn("w:tblW"))
        if tblW is not None:
            tblW.set(qn("w:type"), "pct")
            tblW.set(qn("w:w"), "5000")  # 5000 = 100%
        grid_cols = table._tbl.tblGrid.findall(qn("w:gridCol"))
        col_widths = [int(gc.get(qn("w:w"))) for gc in grid_cols]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    tcW = cell._tc.tcPr.find(qn("w:tcW"))
                    if tcW is not None:
                        tcW.set(qn("w:w"), str(col_widths[i]))
                        tcW.set(qn("w:type"), "dxa")

    # Repair html2docx's post-table output (leaked headings, junk blanks, spacing).
    _fix_table_trailing_content(doc)
    # Drop the leading space html2docx adds after each markdown hard break.
    _strip_space_after_line_breaks(doc)
    # Render `---` thematic breaks as a divider rule (html2docx drops <hr>).
    _render_dividers(doc)

    # Apply the org's configured document styles (fonts/colours). Read at
    # export time so the look follows the org, not whoever wrote the canvas.
    from accounts.models import get_membership
    from core.styles import apply_doc_styles, get_org_styles

    membership = await sync_to_async(get_membership)(request.user)
    apply_doc_styles(doc, get_org_styles(membership.org if membership else None))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = "".join(c for c in canvas.title if c.isalnum() or c in " _-").strip() or "document"
    filename = f"{safe_title}.docx"
    return FileResponse(
        buf,
        as_attachment=True,
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@login_required
@require_POST
def canvas_import(request, thread_id, canvas_id=None):
    """Import a .docx file and convert to markdown canvas content."""
    thread = get_object_or_404(ChatThread, id=thread_id, created_by=request.user)

    uploaded = request.FILES.get("file")
    if not uploaded:
        return HttpResponseBadRequest("No file uploaded.")

    from chat.services import MAX_ATTACHMENT_SIZE, SUPPORTED_DOCX_TYPES, import_docx_to_canvas, set_active_canvas

    # Validate file type
    ct = uploaded.content_type or ""
    is_docx = ct in SUPPORTED_DOCX_TYPES or (uploaded.name and uploaded.name.lower().endswith(".docx"))
    if not is_docx:
        return JsonResponse({"error": "Only .docx files are supported for import."}, status=400)

    # Validate file size
    if uploaded.size > MAX_ATTACHMENT_SIZE:
        return JsonResponse(
            {"error": f"That file is too large (max {MAX_ATTACHMENT_SIZE // (1024 * 1024)} MB)."},
            status=400,
        )

    # Derive the title from the filename (matches import_docx_to_canvas) so the
    # target canvas exists *before* import — embedded images attach to it.
    original_name = uploaded.name or "document"
    title = original_name.rsplit(".", 1)[0][:255] or "Untitled document"

    if canvas_id:
        canvas = get_object_or_404(ChatCanvas, pk=canvas_id, thread=thread)
        canvas.title = title
    else:
        from django.db import IntegrityError
        try:
            canvas = ChatCanvas.objects.get(thread=thread, title=title)
        except ChatCanvas.DoesNotExist:
            try:
                canvas = ChatCanvas.objects.create(thread=thread, title=title, content="")
            except IntegrityError:
                canvas = ChatCanvas.objects.get(thread=thread, title=title)

    _title, content, truncated = import_docx_to_canvas(uploaded, request.user, canvas=canvas)
    canvas.content = content
    canvas.save(update_fields=["title", "content", "updated_at"])

    from chat.services import create_canvas_checkpoint
    cp = create_canvas_checkpoint(canvas, source="import", description="Imported from .docx")
    canvas.accepted_checkpoint = cp
    canvas.save(update_fields=["accepted_checkpoint"])
    set_active_canvas(thread.pk, canvas)

    generated_title = None
    if not thread.title:
        from chat.services import generate_canvas_title

        generated_title = generate_canvas_title(title, content, request.user)
        if generated_title:
            ChatThread.objects.filter(pk=thread.pk).update(title=generated_title)

    resp = {"title": title, "content": content}
    if generated_title:
        resp["thread_title"] = generated_title
    if truncated:
        resp["truncated"] = True
    return JsonResponse(resp)


@login_required
@require_POST
def thread_create(request):
    """Create a new empty thread and return its ID."""
    thread = ChatThread.objects.create(created_by=request.user)
    return JsonResponse({"thread_id": str(thread.id)})


@login_required
@require_POST
def canvas_save_to_data_room(request, thread_id, canvas_id=None):
    """Save canvas content as a markdown document in a data room."""
    thread = get_object_or_404(ChatThread, id=thread_id, created_by=request.user)
    if canvas_id:
        canvas = get_object_or_404(ChatCanvas, pk=canvas_id, thread=thread)
    else:
        canvas = get_object_or_404(ChatCanvas, pk=thread.active_canvas_id, thread=thread)

    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)
    data_room_id = body.get("data_room_id")
    if not data_room_id:
        return JsonResponse({"error": "data_room_id required"}, status=400)

    data_room = get_object_or_404(DataRoom, pk=data_room_id)

    from documents.views import _user_can_modify_data_room

    if not _user_can_modify_data_room(request.user, data_room):
        return JsonResponse({"error": "Access denied"}, status=403)

    from chat.services import save_canvas_to_data_room as save_canvas_to_data_room_service
    from documents.services.sync_scan import scan_version_synchronously

    # Scan synchronously so the button reports the verdict (spinner → checkmark, or the
    # block reason) instead of returning before the async scan finishes. A blocked save
    # is KEPT as an accessible quarantined draft (the work isn't lost); the user sees the
    # reason and can remediate it directly.
    doc, version = save_canvas_to_data_room_service(canvas, data_room, request.user, enqueue=False)
    verdict = scan_version_synchronously(version.id)

    return JsonResponse({
        **verdict.to_http_json(),
        "saved": True,
        "document_id": doc.id,
        "filename": doc.original_filename,
        "data_room_name": data_room.name,
    })


@login_required
@require_http_methods(["GET"])
def skills_for_user(request):
    """JSON API returning the user's available skills."""
    from core.preferences import get_preferences

    prefs = get_preferences(request.user)
    return JsonResponse({
        "skills": [
            {"id": s["id"], "name": s["name"], "description": s.get("description", "")}
            for s in prefs.allowed_skills
        ]
    })


@login_required
@require_POST
def upload_attachments(request, thread_id):
    """Upload file attachments for a chat message.

    Security note: the stored ``content_type`` is client-asserted (the browser's
    reported MIME type, only validated against an allow-list below). It is trusted
    solely for routing the file to the LLM (image vs. pdf vs. text block) and is
    never echoed back to a browser. Attachments are not served back as downloads;
    if a serve-back endpoint is ever added it MUST send ``Content-Disposition:
    attachment`` + ``X-Content-Type-Options: nosniff`` and must NOT reflect this
    stored value as the response ``Content-Type`` (it could be a mislabeled file).
    """
    from chat.services import SUPPORTED_ATTACHMENT_TYPES, SUPPORTED_DOCX_TYPES, max_size_for_content_type

    thread = get_object_or_404(ChatThread, id=thread_id, created_by=request.user)

    files = request.FILES.getlist("files")
    if not files:
        return JsonResponse({"error": "No files uploaded"}, status=400)

    results = []
    for f in files:
        ct = f.content_type
        # Browsers sometimes report .docx as application/octet-stream
        if ct not in SUPPORTED_ATTACHMENT_TYPES:
            if f.name and f.name.lower().endswith(".docx"):
                ct = next(iter(SUPPORTED_DOCX_TYPES))
                f.content_type = ct
            else:
                return JsonResponse(
                    {"error": f"Unsupported file type: {ct}"},
                    status=400,
                )
        max_size = max_size_for_content_type(ct)
        if f.size > max_size:
            return JsonResponse(
                {"error": f"{f.name} is too large (max {max_size // (1024 * 1024)} MB)."},
                status=400,
            )

    for f in files:
        att = ChatAttachment.objects.create(
            thread=thread,
            uploaded_by=request.user,
            file=f,
            original_filename=f.name[:255],
            content_type=f.content_type,
            size_bytes=f.size,
        )
        results.append({
            "id": str(att.id),
            "filename": att.original_filename,
            "content_type": att.content_type,
            "size_bytes": att.size_bytes,
        })

    return JsonResponse({"attachments": results})


@login_required
@require_http_methods(["GET"])
def data_rooms_for_user(request):
    """JSON API returning the user's data rooms for the attach dropdown."""
    rooms = _get_accessible_data_rooms(request.user)
    for r in rooms:
        r["uuid"] = str(r["uuid"])
    return JsonResponse({"data_rooms": rooms})


# Sidebar search: caps keep the response small and bound the icontains scan.
SEARCH_TITLE_LIMIT = 10
SEARCH_CONTENT_LIMIT = 10
SEARCH_CONTENT_MIN_QUERY = 3
SEARCH_CONTENT_SCAN_CAP = 500


def _search_snippet(content, query, before=45, after=65):
    """Short window of ``content`` around the first case-insensitive match of
    ``query``, whitespace-collapsed, with ellipses marking truncation."""
    idx = content.lower().find(query.lower())
    if idx == -1:
        idx = 0
    start = max(0, idx - before)
    end = min(len(content), idx + len(query) + after)
    snippet = " ".join(content[start:end].split())
    if start > 0:
        snippet = "…" + snippet
    if end < len(content):
        snippet = snippet + "…"
    return snippet


@login_required
@require_http_methods(["GET"])
def search_threads(request):
    """JSON API for the sidebar chat search.

    Returns one merged list: threads whose title matches first, then threads
    with a matching user/assistant message (deduped, with a snippet of the
    most recent matching message). Archived threads are included and flagged.
    """
    from chat.models import ChatMessage

    query = request.GET.get("q", "").strip()
    if not query:
        return JsonResponse({"results": []})

    results = []

    title_threads = list(
        ChatThread.objects.filter(created_by=request.user, title__icontains=query)
        .order_by("-updated_at")[:SEARCH_TITLE_LIMIT]
    )
    for t in title_threads:
        results.append({
            "id": str(t.id),
            "title": str(t),
            "emoji": t.emoji,
            "is_archived": t.is_archived,
            "snippet": None,
        })

    # Content matches: skip very short queries (they match everything) and
    # threads already found by title. Hidden messages aren't shown in the UI,
    # so they shouldn't be searchable either.
    if len(query) >= SEARCH_CONTENT_MIN_QUERY:
        title_ids = {t.id for t in title_threads}
        matched = {}  # thread_id -> snippet of most recent matching message
        message_rows = (
            ChatMessage.objects.filter(
                thread__created_by=request.user,
                role__in=[ChatMessage.Role.USER, ChatMessage.Role.ASSISTANT],
                is_hidden_from_user=False,
                content__icontains=query,
            )
            .exclude(thread_id__in=title_ids)
            .order_by("-created_at")
            .values_list("thread_id", "content")[:SEARCH_CONTENT_SCAN_CAP]
        )
        for thread_id, content in message_rows:
            if thread_id in matched:
                continue
            matched[thread_id] = _search_snippet(content, query)
            if len(matched) >= SEARCH_CONTENT_LIMIT:
                break

        if matched:
            content_threads = ChatThread.objects.filter(id__in=matched)
            threads_by_id = {t.id: t for t in content_threads}
            # Preserve message-recency order from the scan above.
            for thread_id, snippet in matched.items():
                t = threads_by_id.get(thread_id)
                if t is None:
                    continue
                results.append({
                    "id": str(t.id),
                    "title": str(t),
                    "emoji": t.emoji,
                    "is_archived": t.is_archived,
                    "snippet": snippet,
                })

    return JsonResponse({"results": results})


# ---------------------------------------------------------------------------
# Loops
# ---------------------------------------------------------------------------

# Loop payload validation, resource linking, and create/edit/pause/restart
# orchestration live in ``chat/loop_service.py`` (shared with the agent tools in
# ``chat/tool_loops.py``). The views below are thin request wrappers around them.


@login_required
@require_http_methods(["GET"])
def loops_list(request):
    """Loops page: active + paused loops with unread indicators and a create modal."""
    import json as _json

    from django.conf import settings as django_settings
    from django.db.models import F

    from chat.loop_service import _loop_form_json
    from chat.models import Loop
    from core.preferences import get_preferences
    from llm.display import get_display_name

    prefs = get_preferences(request.user)

    base = Loop.objects.filter(created_by=request.user).select_related("thread")
    order = [F("last_result_at").desc(nulls_last=True), "-created_at"]
    active_loops = list(base.filter(status=Loop.Status.ACTIVE).order_by(*order))
    paused_loops = list(base.filter(status=Loop.Status.PAUSED).order_by(*order))

    data_rooms = _get_accessible_data_rooms(request.user)
    for r in data_rooms:
        r["uuid"] = str(r["uuid"])
    skills = [
        {"id": str(s["id"]), "name": s["name"], "emoji": s.get("emoji", "")}
        for s in prefs.allowed_skills
    ]
    model_choices = [{"id": m, "display": get_display_name(m)} for m in prefs.allowed_models]
    preferred_chat_model = prefs.feature_models.get("chat", prefs.top_model)
    loops_json = {
        str(loop.id): _loop_form_json(loop, prefs)
        for loop in (active_loops + paused_loops)
    }

    return render(request, "chat/loops_list.html", {
        "active_loops": active_loops,
        "paused_loops": paused_loops,
        "data_rooms": data_rooms,
        "skills": skills,
        "data_rooms_json": _json.dumps([{"id": r["pk"], "name": r["name"]} for r in data_rooms]),
        "skills_json": _json.dumps(skills),
        "model_choices_json": _json.dumps(model_choices),
        "preferred_chat_model": preferred_chat_model,
        "preferred_chat_model_display": get_display_name(preferred_chat_model),
        "loops_json": _json.dumps(loops_json),
        "assistant_name": django_settings.ASSISTANT_NAME,
        # /loop prefill / deep-links
        "open_create": request.GET.get("new") == "1",
        "open_edit_id": request.GET.get("edit", "") if request.GET.get("edit") in loops_json else "",
        "prefill_prompt": request.GET.get("prompt", ""),
        "prefill_interval": request.GET.get("interval", ""),
    })


@login_required
@require_POST
def loop_create(request):
    """Create a loop (and its backing thread) from the modal payload."""
    from django.utils import timezone

    from chat.loop_service import create_loop

    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    loop, errors = create_loop(
        user=request.user, body=body,
        now=timezone.now(), tz_name=timezone.get_current_timezone_name(),
    )
    if errors:
        return JsonResponse({"error": " ".join(errors)}, status=400)

    return JsonResponse({
        "ok": True,
        "loop_id": str(loop.id),
        "thread_id": str(loop.thread_id),
        "max_runs": loop.max_runs,
    })


@login_required
@require_POST
def loop_edit(request, loop_id):
    """Edit an existing loop. Recomputes next_run going forward."""
    from django.utils import timezone

    from chat.loop_service import update_loop
    from chat.models import Loop

    loop = get_object_or_404(Loop.objects.select_related("thread"), id=loop_id, created_by=request.user)
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    loop, errors = update_loop(
        loop=loop, user=request.user, body=body,
        now=timezone.now(), tz_name=timezone.get_current_timezone_name(),
    )
    if errors:
        return JsonResponse({"error": " ".join(errors)}, status=400)

    return JsonResponse({
        "ok": True,
        "loop_id": str(loop.id),
        "max_runs": loop.max_runs,
    })


@login_required
@require_POST
def loop_pause(request, loop_id):
    from chat.loop_service import pause_loop
    from chat.models import Loop

    loop = get_object_or_404(Loop, id=loop_id, created_by=request.user)
    pause_loop(loop)
    return JsonResponse({"ok": True})


@login_required
@require_POST
def loop_restart(request, loop_id):
    """Restart a paused loop: reset the run count and fire on the next tick.

    This is the user-facing revive action. (The agent's "resume where it left
    off" path lives in the loop tools, not here.)
    """
    from django.utils import timezone

    from chat.loop_service import restart_loop
    from chat.models import Loop

    loop = get_object_or_404(Loop, id=loop_id, created_by=request.user)
    restart_loop(loop, timezone.now())
    return JsonResponse({"ok": True})
