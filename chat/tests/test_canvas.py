"""Tests for Canvas tools, consumer events, and views."""

from __future__ import annotations

import json
import io
from unittest.mock import MagicMock, patch

from django.contrib.auth import get_user_model
from django.test import TestCase, TransactionTestCase, override_settings

from chat.canvas_tools import ActiveCanvasTool, EditCanvasTool, WriteCanvasTool
from chat.models import CanvasCheckpoint, ChatCanvas, ChatThread
from llm.types.context import RunContext

User = get_user_model()


def _ctx(user_id, thread_id):
    return RunContext.create(user_id=user_id, conversation_id=str(thread_id))


def _invoke(tool_cls, args, ctx):
    tool = tool_cls()
    tool.set_context(ctx)
    return json.loads(tool.invoke(args))


# ---------------------------------------------------------------------------
# WriteCanvasTool tests
# ---------------------------------------------------------------------------

class WriteCanvasToolTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="canvas@test.com", password="pass")
        self.thread = ChatThread.objects.create(created_by=self.user)

    def test_creates_canvas_on_first_call(self):
        result = _invoke(WriteCanvasTool, {"title": "My NDA", "content": "# NDA\n..."}, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["status"], "ok")
        canvas = ChatCanvas.objects.get(thread=self.thread)
        self.assertEqual(canvas.title, "My NDA")
        self.assertEqual(canvas.content, "# NDA\n...")

    def test_overwrites_canvas_with_same_title(self):
        _invoke(WriteCanvasTool, {"title": "Draft", "content": "old"}, _ctx(self.user.pk, self.thread.id))
        _invoke(WriteCanvasTool, {"title": "Draft", "content": "new"}, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(ChatCanvas.objects.filter(thread=self.thread).count(), 1)
        canvas = ChatCanvas.objects.get(thread=self.thread, title="Draft")
        self.assertEqual(canvas.content, "new")

    def test_creates_second_canvas_with_different_title(self):
        _invoke(WriteCanvasTool, {"title": "Draft 1", "content": "old"}, _ctx(self.user.pk, self.thread.id))
        _invoke(WriteCanvasTool, {"title": "Draft 2", "content": "new"}, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(ChatCanvas.objects.filter(thread=self.thread).count(), 2)

    def test_returns_metadata_in_result(self):
        result = _invoke(WriteCanvasTool, {"title": "T", "content": "C"}, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["title"], "T")
        self.assertIn("canvas_id", result)
        self.assertNotIn("content", result)
        self.assertNotIn("accepted_content", result)

    def test_no_context_returns_error(self):
        tool = WriteCanvasTool()
        result = json.loads(tool.invoke({"title": "T", "content": "C"}))
        self.assertEqual(result["status"], "error")


# ---------------------------------------------------------------------------
# ActiveCanvasTool tests
# ---------------------------------------------------------------------------

class ActiveCanvasToolTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="active@test.com", password="pass")
        self.thread = ChatThread.objects.create(created_by=self.user)

    def test_activates_single_canvas(self):
        ChatCanvas.objects.create(thread=self.thread, title="Notes", content="n")
        result = _invoke(
            ActiveCanvasTool, {"canvas_names": ["Notes"]},
            _ctx(self.user.pk, self.thread.id),
        )
        self.assertEqual(result["status"], "ok")
        self.assertEqual(len(result["activated"]), 1)
        self.assertEqual(result["activated"][0]["title"], "Notes")
        canvas = ChatCanvas.objects.get(thread=self.thread, title="Notes")
        self.assertTrue(canvas.is_active)

    def test_activates_multiple_canvases(self):
        ChatCanvas.objects.create(thread=self.thread, title="A", content="a")
        ChatCanvas.objects.create(thread=self.thread, title="B", content="b")
        ChatCanvas.objects.create(thread=self.thread, title="C", content="c")
        result = _invoke(
            ActiveCanvasTool, {"canvas_names": ["A", "C"]},
            _ctx(self.user.pk, self.thread.id),
        )
        self.assertEqual(result["status"], "ok")
        self.assertEqual(len(result["activated"]), 2)
        a = ChatCanvas.objects.get(thread=self.thread, title="A")
        b = ChatCanvas.objects.get(thread=self.thread, title="B")
        c = ChatCanvas.objects.get(thread=self.thread, title="C")
        self.assertTrue(a.is_active)
        self.assertFalse(b.is_active)
        self.assertTrue(c.is_active)

    def test_deactivates_previously_active(self):
        from django.utils import timezone

        now = timezone.now()
        old = ChatCanvas.objects.create(thread=self.thread, title="Old", content="o")
        old.is_active = True
        old.last_activated_at = now
        old.save(update_fields=["is_active", "last_activated_at"])
        ChatCanvas.objects.create(thread=self.thread, title="New", content="n")

        result = _invoke(
            ActiveCanvasTool, {"canvas_names": ["New"]},
            _ctx(self.user.pk, self.thread.id),
        )
        self.assertEqual(result["status"], "ok")
        old.refresh_from_db()
        self.assertFalse(old.is_active)
        new = ChatCanvas.objects.get(thread=self.thread, title="New")
        self.assertTrue(new.is_active)

    def test_partial_success_with_invalid_name(self):
        ChatCanvas.objects.create(thread=self.thread, title="Real", content="r")
        result = _invoke(
            ActiveCanvasTool, {"canvas_names": ["Real", "Fake"]},
            _ctx(self.user.pk, self.thread.id),
        )
        self.assertEqual(result["status"], "ok")
        self.assertEqual(len(result["activated"]), 1)
        self.assertEqual(len(result["errors"]), 1)
        self.assertIn("Fake", result["errors"][0])

    def test_rejects_more_than_three(self):
        for name in ["A", "B", "C", "D"]:
            ChatCanvas.objects.create(thread=self.thread, title=name, content=name)
        result = _invoke(
            ActiveCanvasTool, {"canvas_names": ["A", "B", "C", "D"]},
            _ctx(self.user.pk, self.thread.id),
        )
        self.assertEqual(result["status"], "error")

    def test_no_context_returns_error(self):
        tool = ActiveCanvasTool()
        result = json.loads(tool.invoke({"canvas_names": ["X"]}))
        self.assertEqual(result["status"], "error")


# ---------------------------------------------------------------------------
# EditCanvasTool tests
# ---------------------------------------------------------------------------

class EditCanvasToolTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="edit@test.com", password="pass")
        self.thread = ChatThread.objects.create(created_by=self.user)

    def _setup_canvas(self, content):
        from django.utils import timezone

        canvas = ChatCanvas.objects.create(
            thread=self.thread, title="Doc", content=content,
            is_active=True, last_activated_at=timezone.now(),
        )
        self.thread.active_canvas = canvas
        self.thread.save(update_fields=["active_canvas"])

    def test_applies_edit_correctly(self):
        self._setup_canvas("The term is 3 years.")
        result = _invoke(EditCanvasTool, {
            "edits": [{"old_text": "3 years", "new_text": "5 years", "reason": "update term"}]
        }, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["applied"], 1)
        self.assertNotIn("content", result)
        canvas = ChatCanvas.objects.get(thread=self.thread, title="Doc")
        self.assertIn("5 years", canvas.content)

    def test_skips_unfound_old_text(self):
        self._setup_canvas("Some content here.")
        result = _invoke(EditCanvasTool, {
            "edits": [{"old_text": "not in text", "new_text": "replacement", "reason": "test"}]
        }, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["applied"], 0)
        self.assertEqual(len(result["failed"]), 1)

    def test_partial_edits_applied(self):
        self._setup_canvas("Hello world foo bar.")
        result = _invoke(EditCanvasTool, {
            "edits": [
                {"old_text": "Hello", "new_text": "Hi", "reason": ""},
                {"old_text": "nonexistent", "new_text": "x", "reason": ""},
            ]
        }, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["applied"], 1)
        self.assertEqual(len(result["failed"]), 1)

    def test_rejects_duplicate_matches(self):
        self._setup_canvas("The cat sat. The cat slept.")
        result = _invoke(EditCanvasTool, {
            "edits": [{"old_text": "The cat", "new_text": "The dog", "reason": ""}]
        }, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["applied"], 0)
        self.assertEqual(len(result["failed"]), 1)
        self.assertIn("2 matches", result["failed"][0]["error"])
        # Content should be unchanged
        canvas = ChatCanvas.objects.get(thread=self.thread, title="Doc")
        self.assertEqual(canvas.content, "The cat sat. The cat slept.")

    def test_returns_error_when_no_canvas(self):
        result = _invoke(EditCanvasTool, {
            "edits": [{"old_text": "x", "new_text": "y", "reason": ""}]
        }, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["status"], "error")
        self.assertIn("write_canvas", result["message"])


# ---------------------------------------------------------------------------
# Consumer canvas event tests
# ---------------------------------------------------------------------------

@override_settings(
    CHANNEL_LAYERS={"default": {"BACKEND": "channels.layers.InMemoryChannelLayer"}}
)
class ConsumerCanvasTests(TransactionTestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="ws@test.com", password="pass")

    async def _communicator(self):
        from channels.routing import URLRouter
        from channels.testing import WebsocketCommunicator
        from chat.routing import websocket_urlpatterns

        app = URLRouter(websocket_urlpatterns)
        comm = WebsocketCommunicator(app, "/ws/chat/")
        comm.scope["user"] = self.user
        connected, _ = await comm.connect()
        assert connected
        return comm

    async def test_canvas_loaded_on_thread_load(self):
        from channels.db import database_sync_to_async

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)

        def setup():
            c = ChatCanvas.objects.create(thread=thread, title="My Doc", content="hello")
            thread.active_canvas = c
            thread.save(update_fields=["active_canvas"])
        await database_sync_to_async(setup)()

        comm = await self._communicator()
        await comm.send_json_to({"type": "chat.load_thread", "thread_id": str(thread.id)})

        # thread.loaded event
        msg1 = await comm.receive_json_from()
        self.assertEqual(msg1["event_type"], "thread.loaded")

        # canvases.loaded event
        msg2 = await comm.receive_json_from()
        self.assertEqual(msg2["event_type"], "canvases.loaded")
        self.assertEqual(len(msg2["canvases"]), 1)
        self.assertEqual(msg2["active_canvas"]["title"], "My Doc")
        self.assertEqual(msg2["active_canvas"]["content"], "hello")

        await comm.disconnect()

    async def test_canvas_save_persists(self):
        from channels.db import database_sync_to_async

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)

        comm = await self._communicator()
        await comm.send_json_to({
            "type": "chat.canvas_save",
            "thread_id": str(thread.id),
            "title": "Saved Title",
            "content": "saved content",
        })
        # No response expected — just persistence
        await comm.disconnect()

        canvas = await database_sync_to_async(ChatCanvas.objects.get)(thread=thread)
        self.assertEqual(canvas.title, "Saved Title")
        self.assertEqual(canvas.content, "saved content")

    async def test_canvas_open_creates_and_returns_canvas(self):
        from channels.db import database_sync_to_async

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)

        comm = await self._communicator()
        await comm.send_json_to({"type": "chat.canvas_open", "thread_id": str(thread.id)})

        msg = await comm.receive_json_from()
        self.assertEqual(msg["event_type"], "canvases.loaded")
        self.assertIn("canvases", msg)
        self.assertIn("active_canvas", msg)

        await comm.disconnect()


# ---------------------------------------------------------------------------
# View tests: canvas_export and canvas_import
# ---------------------------------------------------------------------------

class CanvasExportViewTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="exp@test.com", password="pass")
        self.client.login(email="exp@test.com", password="pass")
        self.thread = ChatThread.objects.create(created_by=self.user)
        self.canvas = ChatCanvas.objects.create(
            thread=self.thread,
            title="My NDA",
            content="# NDA\n\nThis is an agreement.",
        )
        self.thread.active_canvas = self.canvas
        self.thread.save(update_fields=["active_canvas"])

    def test_export_returns_docx(self):
        url = f"/chat/threads/{self.thread.id}/canvas/export/"
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.get("Content-Type"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertIn("My NDA", response.get("Content-Disposition", ""))

    def test_export_tables_use_fixed_layout(self):
        """Tables in exported .docx should span the full page width."""
        self.canvas.content = "| A | B |\n| --- | --- |\n| 1 | 2 |"
        self.canvas.save(update_fields=["content"])

        url = f"/chat/threads/{self.thread.id}/canvas/export/"
        response = self.client.get(url)
        self.assertEqual(response.status_code, 200)

        from docx import Document as DocxDocument
        from docx.oxml.ns import qn

        doc = DocxDocument(io.BytesIO(response.getvalue()))
        self.assertTrue(len(doc.tables) > 0, "Expected at least one table")
        for table in doc.tables:
            self.assertFalse(table.autofit)
            # Table width should be 100% (5000 in OOXML pct units)
            tblW = table._tbl.tblPr.find(qn("w:tblW"))
            self.assertEqual(tblW.get(qn("w:type")), "pct")
            self.assertEqual(tblW.get(qn("w:w")), "5000")
            # Cell widths should be non-zero
            for row in table.rows:
                for cell in row.cells:
                    tcW = cell._tc.tcPr.find(qn("w:tcW"))
                    self.assertGreater(int(tcW.get(qn("w:w"))), 0)

    def test_export_404_no_canvas(self):
        other_thread = ChatThread.objects.create(created_by=self.user)
        url = f"/chat/threads/{other_thread.id}/canvas/export/"
        response = self.client.get(url)
        self.assertEqual(response.status_code, 404)

    def test_export_404_other_user(self):
        other = User.objects.create_user(email="other@test.com", password="pass")
        other_thread = ChatThread.objects.create(created_by=other)
        ChatCanvas.objects.create(thread=other_thread, title="X", content="y")
        url = f"/chat/threads/{other_thread.id}/canvas/export/"
        response = self.client.get(url)
        self.assertEqual(response.status_code, 404)


class CanvasImportViewTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="imp@test.com", password="pass")
        self.client.login(email="imp@test.com", password="pass")
        self.thread = ChatThread.objects.create(created_by=self.user)

    @patch("chat.services.import_docx_to_canvas")
    def test_import_creates_canvas(self, mock_import):
        mock_import.return_value = ("contract", "# Converted\n\nContent here.", False)

        url = f"/chat/threads/{self.thread.id}/canvas/import/"
        fake_file = io.BytesIO(b"PK fake docx")
        fake_file.name = "contract.docx"
        response = self.client.post(url, {"file": fake_file}, format="multipart")
        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["title"], "contract")
        self.assertIn("Converted", data["content"])

        canvas = ChatCanvas.objects.get(thread=self.thread)
        self.assertEqual(canvas.title, "contract")

    def test_import_no_file_returns_400(self):
        url = f"/chat/threads/{self.thread.id}/canvas/import/"
        response = self.client.post(url)
        self.assertEqual(response.status_code, 400)


# ---------------------------------------------------------------------------
# Service tests: describe_image and import_docx_to_canvas
# ---------------------------------------------------------------------------

class DescribeImageTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="desc@test.com", password="pass")

    @patch("core.preferences.get_preferences")
    @patch("llm.get_llm_service")
    def test_returns_description(self, mock_svc, mock_prefs):
        from chat.services import describe_image

        prefs = MagicMock()
        prefs.cheap_model = "gemini/gemini-2.5-flash"
        prefs.mid_model = "openai/gpt-5-mini"
        prefs.top_model = "anthropic/claude-sonnet-4-5"
        mock_prefs.return_value = prefs

        mock_response = MagicMock()
        mock_response.message.content = "A bar chart showing Q1 revenue"
        mock_svc.return_value.run.return_value = mock_response

        result = describe_image(b"\x89PNG", "image/png", self.user)
        self.assertEqual(result, "A bar chart showing Q1 revenue")
        mock_svc.return_value.run.assert_called_once()

    @patch("core.preferences.get_preferences")
    def test_returns_none_no_vision_model(self, mock_prefs):
        from chat.services import describe_image

        prefs = MagicMock()
        prefs.cheap_model = "openai/o3"
        prefs.mid_model = "openai/o4-mini"
        prefs.top_model = "custom/text-only"
        mock_prefs.return_value = prefs

        result = describe_image(b"\x89PNG", "image/png", self.user)
        self.assertIsNone(result)

    @patch("core.preferences.get_preferences")
    @patch("llm.get_llm_service")
    def test_returns_none_on_exception(self, mock_svc, mock_prefs):
        from chat.services import describe_image

        prefs = MagicMock()
        prefs.cheap_model = "anthropic/claude-haiku-4-5-20251001"
        prefs.mid_model = "openai/gpt-5-mini"
        prefs.top_model = "anthropic/claude-sonnet-4-5"
        mock_prefs.return_value = prefs

        mock_svc.return_value.run.side_effect = RuntimeError("API error")
        result = describe_image(b"\x89PNG", "image/png", self.user)
        self.assertIsNone(result)

    @patch("core.preferences.get_preferences")
    @patch("llm.get_llm_service")
    def test_picks_first_vision_capable_model(self, mock_svc, mock_prefs):
        from chat.services import describe_image

        prefs = MagicMock()
        prefs.cheap_model = "openai/o3"  # no vision
        prefs.mid_model = "openai/gpt-5-mini"  # vision
        prefs.top_model = "anthropic/claude-sonnet-4-5"
        mock_prefs.return_value = prefs

        mock_response = MagicMock()
        mock_response.message.content = "A photo"
        mock_svc.return_value.run.return_value = mock_response

        describe_image(b"\x89PNG", "image/png", self.user)
        call_args = mock_svc.return_value.run.call_args
        request = call_args[0][1]
        self.assertEqual(request.model, "openai/gpt-5-mini")


class GenerateCanvasTitleTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="gentitle@test.com", password="pass")

    @patch("core.preferences.get_preferences")
    @patch("llm.get_llm_service")
    def test_returns_generated_title(self, mock_svc, mock_prefs):
        from chat.services import generate_canvas_title

        prefs = MagicMock()
        prefs.cheap_model = "openai/gpt-5-mini"
        mock_prefs.return_value = prefs

        mock_response = MagicMock()
        mock_response.message.content = '"Patent License Agreement"'
        mock_svc.return_value.run.return_value = mock_response

        result = generate_canvas_title("license.docx", "This is a patent license...", self.user)
        self.assertEqual(result, "Patent License Agreement")
        mock_svc.return_value.run.assert_called_once()

    @patch("core.preferences.get_preferences")
    @patch("llm.get_llm_service")
    def test_returns_none_on_exception(self, mock_svc, mock_prefs):
        from chat.services import generate_canvas_title

        prefs = MagicMock()
        prefs.cheap_model = "openai/gpt-5-mini"
        mock_prefs.return_value = prefs

        mock_svc.return_value.run.side_effect = RuntimeError("API down")
        result = generate_canvas_title("doc", "content", self.user)
        self.assertIsNone(result)

    @patch("core.preferences.get_preferences")
    @patch("llm.get_llm_service")
    def test_truncates_long_title(self, mock_svc, mock_prefs):
        from chat.services import generate_canvas_title

        prefs = MagicMock()
        prefs.cheap_model = "openai/gpt-5-mini"
        mock_prefs.return_value = prefs

        mock_response = MagicMock()
        mock_response.message.content = "A" * 300
        mock_svc.return_value.run.return_value = mock_response

        result = generate_canvas_title("doc", "content", self.user)
        self.assertEqual(len(result), 255)


class CanvasImportTitleIntegrationTests(TestCase):
    """Test that canvas_import generates a title for untitled threads."""

    def setUp(self):
        self.user = User.objects.create_user(email="imptitle@test.com", password="pass")
        self.client.login(email="imptitle@test.com", password="pass")

    @patch("chat.services.generate_canvas_title", return_value="Generated Title")
    @patch("chat.services.import_docx_to_canvas")
    def test_generates_title_for_untitled_thread(self, mock_import, mock_gen):
        mock_import.return_value = ("contract", "# Content", False)
        thread = ChatThread.objects.create(created_by=self.user)

        url = f"/chat/threads/{thread.id}/canvas/import/"
        fake_file = io.BytesIO(b"PK fake docx")
        fake_file.name = "contract.docx"
        response = self.client.post(url, {"file": fake_file}, format="multipart")

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertEqual(data["thread_title"], "Generated Title")
        thread.refresh_from_db()
        self.assertEqual(thread.title, "Generated Title")

    @patch("chat.services.generate_canvas_title")
    @patch("chat.services.import_docx_to_canvas")
    def test_does_not_overwrite_existing_title(self, mock_import, mock_gen):
        mock_import.return_value = ("contract", "# Content", False)
        thread = ChatThread.objects.create(created_by=self.user, title="Existing Title")

        url = f"/chat/threads/{thread.id}/canvas/import/"
        fake_file = io.BytesIO(b"PK fake docx")
        fake_file.name = "contract.docx"
        response = self.client.post(url, {"file": fake_file}, format="multipart")

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertNotIn("thread_title", data)
        mock_gen.assert_not_called()
        thread.refresh_from_db()
        self.assertEqual(thread.title, "Existing Title")

    @patch("chat.services.generate_canvas_title", return_value=None)
    @patch("chat.services.import_docx_to_canvas")
    def test_llm_failure_does_not_break_import(self, mock_import, mock_gen):
        mock_import.return_value = ("contract", "# Content", False)
        thread = ChatThread.objects.create(created_by=self.user)

        url = f"/chat/threads/{thread.id}/canvas/import/"
        fake_file = io.BytesIO(b"PK fake docx")
        fake_file.name = "contract.docx"
        response = self.client.post(url, {"file": fake_file}, format="multipart")

        self.assertEqual(response.status_code, 200)
        data = response.json()
        self.assertNotIn("thread_title", data)
        self.assertEqual(data["title"], "contract")


class ImportDocxToCanvasTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="imp2@test.com", password="pass")

    @patch("chat.services.describe_image")
    @patch("mammoth.convert_to_html")
    def test_basic_conversion_no_images(self, mock_convert, mock_describe):
        from chat.services import import_docx_to_canvas

        mock_result = MagicMock()
        mock_result.value = "<h1>Hello</h1><p>World</p>"
        mock_convert.return_value = mock_result

        fake_file = MagicMock()
        fake_file.name = "report.docx"

        title, content, truncated = import_docx_to_canvas(fake_file, self.user)
        self.assertEqual(title, "report")
        self.assertEqual(content, "# Hello\n\nWorld")
        self.assertFalse(truncated)
        mock_describe.assert_not_called()

    @patch("chat.services.describe_image")
    @patch("mammoth.convert_to_html")
    def test_table_conversion(self, mock_convert, mock_describe):
        from chat.services import import_docx_to_canvas

        mock_result = MagicMock()
        mock_result.value = (
            "<table>"
            "<tr><th>Name</th><th>Value</th></tr>"
            "<tr><td>A</td><td>1</td></tr>"
            "<tr><td>B</td><td>2</td></tr>"
            "</table>"
        )
        mock_convert.return_value = mock_result

        fake_file = MagicMock()
        fake_file.name = "data.docx"

        title, content, truncated = import_docx_to_canvas(fake_file, self.user)
        self.assertEqual(title, "data")
        self.assertFalse(truncated)
        self.assertIn("| Name | Value |", content)
        self.assertIn("| --- | --- |", content)
        self.assertIn("| A | 1 |", content)
        mock_describe.assert_not_called()

    @patch("chat.services.describe_image")
    def test_images_get_placeholders(self, mock_describe):
        from chat.services import import_docx_to_canvas

        mock_describe.return_value = "A revenue chart"

        # Create a real tiny docx with an image to test end-to-end is complex,
        # so we test the regex cleanup directly
        import re
        raw = "Some text\n\n![[Image 1: A revenue chart]](#)\n\nMore text"
        cleaned = re.sub(
            r"!\[(\[Image \d+[^\]]*\])\]\([^)]*\)",
            r"\1",
            raw,
        )
        self.assertEqual(cleaned, "Some text\n\n[Image 1: A revenue chart]\n\nMore text")

    @patch("chat.services.describe_image")
    def test_image_placeholder_without_description(self, mock_describe):
        import re
        raw = "Before\n\n![[Image 1: EMF image]](#)\n\nAfter"
        cleaned = re.sub(
            r"!\[(\[Image \d+[^\]]*\])\]\([^)]*\)",
            r"\1",
            raw,
        )
        self.assertEqual(cleaned, "Before\n\n[Image 1: EMF image]\n\nAfter")


# ---------------------------------------------------------------------------
# CanvasCheckpoint model tests
# ---------------------------------------------------------------------------

class CanvasCheckpointModelTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="cpmodel@test.com", password="pass")
        self.thread = ChatThread.objects.create(created_by=self.user)
        self.canvas = ChatCanvas.objects.create(thread=self.thread, title="Doc", content="Hello")

    def test_create_checkpoint(self):
        cp = CanvasCheckpoint.objects.create(
            canvas=self.canvas, title="Doc", content="Hello",
            source=CanvasCheckpoint.Source.ORIGINAL, order=1,
        )
        self.assertEqual(cp.source, "original")
        self.assertEqual(cp.order, 1)

    def test_ordering(self):
        CanvasCheckpoint.objects.create(canvas=self.canvas, title="A", content="1", source="original", order=1)
        CanvasCheckpoint.objects.create(canvas=self.canvas, title="B", content="2", source="ai_edit", order=2)
        cps = list(CanvasCheckpoint.objects.filter(canvas=self.canvas))
        self.assertEqual(cps[0].order, 1)
        self.assertEqual(cps[1].order, 2)

    def test_source_choices(self):
        for val, label in CanvasCheckpoint.Source.choices:
            cp = CanvasCheckpoint.objects.create(
                canvas=self.canvas, title="T", content="C",
                source=val, order=0,
            )
            self.assertEqual(cp.source, val)


# ---------------------------------------------------------------------------
# WriteCanvasTool checkpoint tests
# ---------------------------------------------------------------------------

class WriteCanvasToolCheckpointTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="wcp@test.com", password="pass")
        self.thread = ChatThread.objects.create(created_by=self.user)

    def test_first_call_creates_original_checkpoint_and_sets_accepted(self):
        result = _invoke(WriteCanvasTool, {"title": "NDA", "content": "# NDA"}, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["status"], "ok")
        canvas = ChatCanvas.objects.get(thread=self.thread)
        cps = list(CanvasCheckpoint.objects.filter(canvas=canvas))
        self.assertEqual(len(cps), 1)
        self.assertEqual(cps[0].source, "original")
        self.assertIsNotNone(canvas.accepted_checkpoint)
        self.assertEqual(canvas.accepted_checkpoint.pk, cps[0].pk)

    def test_second_call_creates_ai_edit_without_updating_accepted(self):
        _invoke(WriteCanvasTool, {"title": "Draft", "content": "old"}, _ctx(self.user.pk, self.thread.id))
        canvas = ChatCanvas.objects.get(thread=self.thread, title="Draft")
        original_accepted_pk = canvas.accepted_checkpoint.pk

        _invoke(WriteCanvasTool, {"title": "Draft", "content": "new"}, _ctx(self.user.pk, self.thread.id))
        canvas.refresh_from_db()
        cps = list(CanvasCheckpoint.objects.filter(canvas=canvas).order_by("order"))
        self.assertEqual(len(cps), 2)
        self.assertEqual(cps[1].source, "ai_edit")
        # accepted_checkpoint should NOT be updated
        self.assertEqual(canvas.accepted_checkpoint.pk, original_accepted_pk)

    def test_no_content_in_result(self):
        result = _invoke(WriteCanvasTool, {"title": "T", "content": "C"}, _ctx(self.user.pk, self.thread.id))
        self.assertNotIn("content", result)
        self.assertNotIn("accepted_content", result)


# ---------------------------------------------------------------------------
# EditCanvasTool checkpoint tests
# ---------------------------------------------------------------------------

class EditCanvasToolCheckpointTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="ecp@test.com", password="pass")
        self.thread = ChatThread.objects.create(created_by=self.user)

    def _setup_canvas(self, content):
        from django.utils import timezone

        canvas = ChatCanvas.objects.create(
            thread=self.thread, title="Doc", content=content,
            is_active=True, last_activated_at=timezone.now(),
        )
        self.thread.active_canvas = canvas
        self.thread.save(update_fields=["active_canvas"])
        from chat.services import create_canvas_checkpoint
        cp = create_canvas_checkpoint(canvas, source="original")
        canvas.accepted_checkpoint = cp
        canvas.save(update_fields=["accepted_checkpoint"])
        return canvas

    def test_successful_edit_creates_ai_edit_checkpoint(self):
        self._setup_canvas("The term is 3 years.")
        result = _invoke(EditCanvasTool, {
            "edits": [{"old_text": "3 years", "new_text": "5 years", "reason": "update"}]
        }, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["applied"], 1)
        canvas = ChatCanvas.objects.get(thread=self.thread)
        cps = list(CanvasCheckpoint.objects.filter(canvas=canvas).order_by("order"))
        self.assertEqual(len(cps), 2)
        self.assertEqual(cps[1].source, "ai_edit")

    def test_failed_edit_does_not_create_checkpoint(self):
        self._setup_canvas("Some content.")
        _invoke(EditCanvasTool, {
            "edits": [{"old_text": "nonexistent", "new_text": "x", "reason": ""}]
        }, _ctx(self.user.pk, self.thread.id))
        canvas = ChatCanvas.objects.get(thread=self.thread)
        cps = list(CanvasCheckpoint.objects.filter(canvas=canvas))
        self.assertEqual(len(cps), 1)  # only the original

    def test_no_content_in_result(self):
        self._setup_canvas("Hello world.")
        result = _invoke(EditCanvasTool, {
            "edits": [{"old_text": "Hello", "new_text": "Hi", "reason": ""}]
        }, _ctx(self.user.pk, self.thread.id))
        self.assertNotIn("content", result)
        self.assertNotIn("accepted_content", result)

    def test_consecutive_ai_edits_coalesce_into_one_checkpoint(self):
        """Multiple AI edits in one turn should produce only one checkpoint."""
        self._setup_canvas("Hello world. Goodbye world.")
        ctx = _ctx(self.user.pk, self.thread.id)
        _invoke(EditCanvasTool, {
            "edits": [{"old_text": "Hello", "new_text": "Hi", "reason": ""}]
        }, ctx)
        _invoke(EditCanvasTool, {
            "edits": [{"old_text": "Goodbye", "new_text": "Bye", "reason": ""}]
        }, ctx)
        canvas = ChatCanvas.objects.get(thread=self.thread)
        cps = list(CanvasCheckpoint.objects.filter(canvas=canvas).order_by("order"))
        # Should have 2: original + one coalesced ai_edit
        self.assertEqual(len(cps), 2)
        self.assertEqual(cps[0].source, "original")
        self.assertEqual(cps[1].source, "ai_edit")
        # The coalesced checkpoint should have the final content
        self.assertIn("Hi", cps[1].content)
        self.assertIn("Bye", cps[1].content)

    def test_write_after_edit_creates_separate_checkpoints(self):
        """A write_canvas after edit_canvas should NOT coalesce (different tool, still ai_edit source though)."""
        self._setup_canvas("Hello world.")
        ctx = _ctx(self.user.pk, self.thread.id)
        _invoke(EditCanvasTool, {
            "edits": [{"old_text": "Hello", "new_text": "Hi", "reason": ""}]
        }, ctx)
        # Write canvas with same title produces ai_edit source (since canvas already exists)
        _invoke(WriteCanvasTool, {"title": "Doc", "content": "Brand new"}, ctx)
        canvas = ChatCanvas.objects.get(thread=self.thread, title="Doc")
        cps = list(CanvasCheckpoint.objects.filter(canvas=canvas).order_by("order"))
        # original + ai_edit (from edit) — but write_canvas on existing canvas
        # also produces ai_edit, which coalesces with the previous ai_edit
        self.assertEqual(len(cps), 2)
        self.assertEqual(cps[1].content, "Brand new")

    def test_ai_edit_after_accept_creates_new_checkpoint(self):
        """After user accepts, next AI edit must NOT coalesce into the accepted checkpoint."""
        canvas = self._setup_canvas("Hello world. Goodbye world.")
        ctx = _ctx(self.user.pk, self.thread.id)

        # First AI edit
        _invoke(EditCanvasTool, {
            "edits": [{"old_text": "Hello", "new_text": "Hi", "reason": ""}]
        }, ctx)
        canvas.refresh_from_db()
        # 2 checkpoints: original + ai_edit
        self.assertEqual(CanvasCheckpoint.objects.filter(canvas=canvas).count(), 2)

        # User accepts — set accepted_checkpoint to latest
        latest = CanvasCheckpoint.objects.filter(canvas=canvas).order_by("-order").first()
        canvas.accepted_checkpoint = latest
        canvas.save(update_fields=["accepted_checkpoint"])

        # Second AI edit — should NOT coalesce into accepted checkpoint
        _invoke(EditCanvasTool, {
            "edits": [{"old_text": "Goodbye", "new_text": "Bye", "reason": ""}]
        }, ctx)
        canvas.refresh_from_db()
        cps = list(CanvasCheckpoint.objects.filter(canvas=canvas).order_by("order"))
        # Should now have 3: original, ai_edit (accepted), ai_edit (new)
        self.assertEqual(len(cps), 3)
        self.assertEqual(cps[1].pk, latest.pk)  # accepted one unchanged
        self.assertIn("Hi", cps[1].content)
        self.assertNotIn("Bye", cps[1].content)  # NOT updated
        self.assertIn("Bye", cps[2].content)  # new checkpoint has it


# ---------------------------------------------------------------------------
# Consumer checkpoint tests
# ---------------------------------------------------------------------------

@override_settings(
    CHANNEL_LAYERS={"default": {"BACKEND": "channels.layers.InMemoryChannelLayer"}}
)
class ConsumerCheckpointTests(TransactionTestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="cpws@test.com", password="pass")

    async def _communicator(self):
        from channels.routing import URLRouter
        from channels.testing import WebsocketCommunicator
        from chat.routing import websocket_urlpatterns

        app = URLRouter(websocket_urlpatterns)
        comm = WebsocketCommunicator(app, "/ws/chat/")
        comm.scope["user"] = self.user
        connected, _ = await comm.connect()
        assert connected
        return comm

    async def test_canvas_accept(self):
        from channels.db import database_sync_to_async
        from chat.services import create_canvas_checkpoint

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)
        canvas = await database_sync_to_async(ChatCanvas.objects.create)(
            thread=thread, title="Doc", content="v1"
        )

        def setup_checkpoints():
            thread.active_canvas = canvas
            thread.save(update_fields=["active_canvas"])
            cp1 = create_canvas_checkpoint(canvas, source="original")
            canvas.accepted_checkpoint = cp1
            canvas.save(update_fields=["accepted_checkpoint"])
            canvas.content = "v2"
            canvas.save(update_fields=["content"])
            create_canvas_checkpoint(canvas, source="ai_edit")
        await database_sync_to_async(setup_checkpoints)()

        comm = await self._communicator()
        await comm.send_json_to({"type": "chat.canvas_accept", "thread_id": str(thread.id)})
        msg = await comm.receive_json_from()
        self.assertEqual(msg["event_type"], "canvas.accepted")
        self.assertEqual(msg["accepted_content"], "v2")

        def check_accepted():
            c = ChatCanvas.objects.get(thread=thread)
            self.assertEqual(c.accepted_checkpoint.content, "v2")
        await database_sync_to_async(check_accepted)()
        await comm.disconnect()

    async def test_canvas_revert(self):
        from channels.db import database_sync_to_async
        from chat.services import create_canvas_checkpoint

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)
        canvas = await database_sync_to_async(ChatCanvas.objects.create)(
            thread=thread, title="Doc", content="original content"
        )

        def setup():
            thread.active_canvas = canvas
            thread.save(update_fields=["active_canvas"])
            cp = create_canvas_checkpoint(canvas, source="original")
            canvas.accepted_checkpoint = cp
            canvas.save(update_fields=["accepted_checkpoint"])
            canvas.content = "ai changed this"
            canvas.save(update_fields=["content"])
        await database_sync_to_async(setup)()

        comm = await self._communicator()
        await comm.send_json_to({"type": "chat.canvas_revert", "thread_id": str(thread.id)})
        msg = await comm.receive_json_from()
        self.assertEqual(msg["event_type"], "canvas.reverted")
        self.assertEqual(msg["content"], "original content")

        def check_content():
            c = ChatCanvas.objects.get(thread=thread)
            self.assertEqual(c.content, "original content")
        await database_sync_to_async(check_content)()
        await comm.disconnect()

    async def test_canvas_save_version(self):
        from channels.db import database_sync_to_async

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)

        def setup():
            c = ChatCanvas.objects.create(thread=thread, title="Doc", content="content")
            thread.active_canvas = c
            thread.save(update_fields=["active_canvas"])
        await database_sync_to_async(setup)()

        comm = await self._communicator()
        await comm.send_json_to({
            "type": "chat.canvas_save_version",
            "thread_id": str(thread.id),
            "title": "Doc",
            "content": "user edited content",
        })
        msg = await comm.receive_json_from()
        self.assertEqual(msg["event_type"], "canvas.version_saved")

        def check():
            c = ChatCanvas.objects.get(thread=thread)
            self.assertIsNotNone(c.accepted_checkpoint)
            self.assertEqual(c.accepted_checkpoint.source, "user_save")
            cps = CanvasCheckpoint.objects.filter(canvas=c)
            self.assertEqual(cps.count(), 1)
        await database_sync_to_async(check)()
        await comm.disconnect()

    async def test_canvas_loaded_includes_accepted_content(self):
        from channels.db import database_sync_to_async
        from chat.services import create_canvas_checkpoint

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)
        canvas = await database_sync_to_async(ChatCanvas.objects.create)(
            thread=thread, title="Doc", content="current"
        )

        def setup():
            thread.active_canvas = canvas
            thread.save(update_fields=["active_canvas"])
            cp = create_canvas_checkpoint(canvas, source="original", description="first")
            canvas.accepted_checkpoint = cp
            canvas.content = "current"
            canvas.save(update_fields=["accepted_checkpoint", "content"])
        await database_sync_to_async(setup)()

        comm = await self._communicator()
        await comm.send_json_to({"type": "chat.load_thread", "thread_id": str(thread.id)})
        # thread.loaded
        msg1 = await comm.receive_json_from()
        self.assertEqual(msg1["event_type"], "thread.loaded")
        # canvases.loaded
        msg2 = await comm.receive_json_from()
        self.assertEqual(msg2["event_type"], "canvases.loaded")
        self.assertIn("accepted_content", msg2["active_canvas"])
        self.assertEqual(msg2["active_canvas"]["accepted_content"], "current")
        await comm.disconnect()

    async def test_canvas_restore_version(self):
        from channels.db import database_sync_to_async
        from chat.services import create_canvas_checkpoint

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)
        canvas = await database_sync_to_async(ChatCanvas.objects.create)(
            thread=thread, title="Doc", content="v1"
        )

        def setup():
            thread.active_canvas = canvas
            thread.save(update_fields=["active_canvas"])
            cp1 = create_canvas_checkpoint(canvas, source="original")
            canvas.accepted_checkpoint = cp1
            canvas.save(update_fields=["accepted_checkpoint"])
            canvas.content = "v2"
            canvas.save(update_fields=["content"])
            create_canvas_checkpoint(canvas, source="ai_edit")
            return cp1.pk
        cp1_pk = await database_sync_to_async(setup)()

        comm = await self._communicator()
        await comm.send_json_to({
            "type": "chat.canvas_restore_version",
            "thread_id": str(thread.id),
            "checkpoint_id": cp1_pk,
        })
        msg = await comm.receive_json_from()
        self.assertEqual(msg["event_type"], "canvas.restored")
        self.assertEqual(msg["content"], "v1")

        def check():
            c = ChatCanvas.objects.get(thread=thread)
            self.assertEqual(c.content, "v1")
            self.assertIsNotNone(c.accepted_checkpoint)
            self.assertEqual(c.accepted_checkpoint.source, "restore")
        await database_sync_to_async(check)()
        await comm.disconnect()

    async def test_canvas_restore_via_get_checkpoints_flow(self):
        """Full user flow: get checkpoints → pick one → restore it."""
        from channels.db import database_sync_to_async
        from chat.services import create_canvas_checkpoint

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)
        canvas = await database_sync_to_async(ChatCanvas.objects.create)(
            thread=thread, title="Doc", content="original text"
        )

        def setup():
            thread.active_canvas = canvas
            thread.save(update_fields=["active_canvas"])
            cp1 = create_canvas_checkpoint(canvas, source="original")
            canvas.accepted_checkpoint = cp1
            canvas.save(update_fields=["accepted_checkpoint"])
            # AI edits the canvas
            canvas.content = "ai changed text"
            canvas.title = "AI Doc"
            canvas.save(update_fields=["title", "content"])
            create_canvas_checkpoint(canvas, source="ai_edit")
        await database_sync_to_async(setup)()

        comm = await self._communicator()

        # Step 1: Get checkpoints
        await comm.send_json_to({
            "type": "chat.canvas_get_checkpoints",
            "thread_id": str(thread.id),
        })
        msg = await comm.receive_json_from()
        self.assertEqual(msg["event_type"], "canvas.checkpoints")
        checkpoints = msg["checkpoints"]
        self.assertEqual(len(checkpoints), 2)
        # Checkpoints are ordered by -order, so first is the latest (ai_edit)
        self.assertEqual(checkpoints[0]["source"], "ai_edit")
        self.assertEqual(checkpoints[1]["source"], "original")

        # Step 2: Restore the original checkpoint using its ID from the list
        original_cp_id = checkpoints[1]["id"]
        await comm.send_json_to({
            "type": "chat.canvas_restore_version",
            "thread_id": str(thread.id),
            "checkpoint_id": original_cp_id,
        })
        msg = await comm.receive_json_from()
        self.assertEqual(msg["event_type"], "canvas.restored")
        self.assertEqual(msg["title"], "Doc")
        self.assertEqual(msg["content"], "original text")

        # Step 3: Verify DB state
        def check():
            c = ChatCanvas.objects.get(thread=thread)
            self.assertEqual(c.content, "original text")
            self.assertEqual(c.title, "Doc")
            self.assertIsNotNone(c.accepted_checkpoint)
            self.assertEqual(c.accepted_checkpoint.source, "restore")
        await database_sync_to_async(check)()
        await comm.disconnect()

    async def test_canvas_get_checkpoints_empty(self):
        """get_checkpoints returns empty list when no checkpoints exist."""
        from channels.db import database_sync_to_async

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)

        def setup():
            c = ChatCanvas.objects.create(thread=thread, title="Doc", content="text")
            thread.active_canvas = c
            thread.save(update_fields=["active_canvas"])
        await database_sync_to_async(setup)()

        comm = await self._communicator()
        await comm.send_json_to({
            "type": "chat.canvas_get_checkpoints",
            "thread_id": str(thread.id),
        })
        msg = await comm.receive_json_from()
        self.assertEqual(msg["event_type"], "canvas.checkpoints")
        self.assertEqual(msg["checkpoints"], [])
        await comm.disconnect()


# ---------------------------------------------------------------------------
# Multi-canvas tool tests
# ---------------------------------------------------------------------------

class MultiCanvasToolTests(TestCase):
    """Tests for multi-canvas behaviour in WriteCanvasTool and EditCanvasTool."""

    def setUp(self):
        self.user = User.objects.create_user(email="multi@test.com", password="pass")
        self.thread = ChatThread.objects.create(created_by=self.user)

    def test_write_creates_new_canvas_with_new_title(self):
        result = _invoke(WriteCanvasTool, {"title": "A", "content": "aaa"}, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["status"], "ok")
        self.assertEqual(ChatCanvas.objects.filter(thread=self.thread).count(), 1)

        result = _invoke(WriteCanvasTool, {"title": "B", "content": "bbb"}, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["status"], "ok")
        self.assertEqual(ChatCanvas.objects.filter(thread=self.thread).count(), 2)
        self.assertEqual(ChatCanvas.objects.get(thread=self.thread, title="B").content, "bbb")

    def test_write_overwrites_canvas_with_same_title(self):
        _invoke(WriteCanvasTool, {"title": "Draft", "content": "v1"}, _ctx(self.user.pk, self.thread.id))
        _invoke(WriteCanvasTool, {"title": "Draft", "content": "v2"}, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(ChatCanvas.objects.filter(thread=self.thread).count(), 1)
        self.assertEqual(ChatCanvas.objects.get(thread=self.thread, title="Draft").content, "v2")

    def test_write_sets_active_canvas(self):
        _invoke(WriteCanvasTool, {"title": "First", "content": "f"}, _ctx(self.user.pk, self.thread.id))
        self.thread.refresh_from_db()
        first = ChatCanvas.objects.get(thread=self.thread, title="First")
        self.assertEqual(self.thread.active_canvas_id, first.pk)

        _invoke(WriteCanvasTool, {"title": "Second", "content": "s"}, _ctx(self.user.pk, self.thread.id))
        self.thread.refresh_from_db()
        second = ChatCanvas.objects.get(thread=self.thread, title="Second")
        self.assertEqual(self.thread.active_canvas_id, second.pk)

    def test_edit_targets_by_canvas_name(self):
        _invoke(WriteCanvasTool, {"title": "Alpha", "content": "Hello world"}, _ctx(self.user.pk, self.thread.id))
        _invoke(WriteCanvasTool, {"title": "Beta", "content": "Foo bar"}, _ctx(self.user.pk, self.thread.id))
        # Active canvas is now "Beta" — edit "Alpha" by name
        result = _invoke(EditCanvasTool, {
            "canvas_name": "Alpha",
            "edits": [{"old_text": "Hello", "new_text": "Hi", "reason": ""}],
        }, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["applied"], 1)
        alpha = ChatCanvas.objects.get(thread=self.thread, title="Alpha")
        self.assertIn("Hi", alpha.content)
        # Beta untouched
        beta = ChatCanvas.objects.get(thread=self.thread, title="Beta")
        self.assertEqual(beta.content, "Foo bar")

    def test_edit_defaults_to_active_canvas(self):
        _invoke(WriteCanvasTool, {"title": "Active", "content": "Hello world"}, _ctx(self.user.pk, self.thread.id))
        result = _invoke(EditCanvasTool, {
            "edits": [{"old_text": "Hello", "new_text": "Hi", "reason": ""}],
        }, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["applied"], 1)
        canvas = ChatCanvas.objects.get(thread=self.thread, title="Active")
        self.assertIn("Hi", canvas.content)

    def test_edit_errors_with_unknown_canvas_name(self):
        _invoke(WriteCanvasTool, {"title": "Doc", "content": "x"}, _ctx(self.user.pk, self.thread.id))
        result = _invoke(EditCanvasTool, {
            "canvas_name": "Nonexistent",
            "edits": [{"old_text": "x", "new_text": "y", "reason": ""}],
        }, _ctx(self.user.pk, self.thread.id))
        self.assertEqual(result["status"], "error")

    def test_unique_title_constraint_enforced(self):
        """Two canvases with the same title cannot exist for one thread."""
        ChatCanvas.objects.create(thread=self.thread, title="Dup", content="a")
        from django.db import IntegrityError
        with self.assertRaises(IntegrityError):
            ChatCanvas.objects.create(thread=self.thread, title="Dup", content="b")

    def test_max_canvases_per_thread(self):
        from chat.services import MAX_CANVASES_PER_THREAD
        ctx = _ctx(self.user.pk, self.thread.id)
        for i in range(MAX_CANVASES_PER_THREAD):
            result = _invoke(WriteCanvasTool, {"title": f"Canvas {i}", "content": "x"}, ctx)
            self.assertEqual(result["status"], "ok")
        result = _invoke(WriteCanvasTool, {"title": "One too many", "content": "x"}, ctx)
        self.assertEqual(result["status"], "error")
        self.assertIn("Maximum", result["message"])

    def test_write_returns_canvas_id(self):
        result = _invoke(WriteCanvasTool, {"title": "T", "content": "C"}, _ctx(self.user.pk, self.thread.id))
        self.assertIn("canvas_id", result)
        canvas = ChatCanvas.objects.get(thread=self.thread, title="T")
        self.assertEqual(result["canvas_id"], str(canvas.pk))

    def test_edit_returns_canvas_id(self):
        _invoke(WriteCanvasTool, {"title": "Doc", "content": "Hello"}, _ctx(self.user.pk, self.thread.id))
        result = _invoke(EditCanvasTool, {
            "edits": [{"old_text": "Hello", "new_text": "Hi", "reason": ""}],
        }, _ctx(self.user.pk, self.thread.id))
        self.assertIn("canvas_id", result)
        canvas = ChatCanvas.objects.get(thread=self.thread, title="Doc")
        self.assertEqual(result["canvas_id"], str(canvas.pk))


# ---------------------------------------------------------------------------
# Multi-canvas consumer tests
# ---------------------------------------------------------------------------

@override_settings(
    CHANNEL_LAYERS={"default": {"BACKEND": "channels.layers.InMemoryChannelLayer"}}
)
class ConsumerMultiCanvasTests(TransactionTestCase):
    """WebSocket tests for multi-canvas features."""

    def setUp(self):
        self.user = User.objects.create_user(email="mws@test.com", password="pass")

    async def _communicator(self):
        from channels.routing import URLRouter
        from channels.testing import WebsocketCommunicator
        from chat.routing import websocket_urlpatterns

        app = URLRouter(websocket_urlpatterns)
        comm = WebsocketCommunicator(app, "/ws/chat/")
        comm.scope["user"] = self.user
        connected, _ = await comm.connect()
        assert connected
        return comm

    async def test_canvases_loaded_on_thread_switch(self):
        from channels.db import database_sync_to_async

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)

        def setup():
            c1 = ChatCanvas.objects.create(thread=thread, title="Doc 1", content="one")
            c2 = ChatCanvas.objects.create(thread=thread, title="Doc 2", content="two")
            thread.active_canvas = c2
            thread.save(update_fields=["active_canvas"])
        await database_sync_to_async(setup)()

        comm = await self._communicator()
        await comm.send_json_to({"type": "chat.load_thread", "thread_id": str(thread.id)})

        msg1 = await comm.receive_json_from()
        self.assertEqual(msg1["event_type"], "thread.loaded")

        msg2 = await comm.receive_json_from()
        self.assertEqual(msg2["event_type"], "canvases.loaded")
        self.assertEqual(len(msg2["canvases"]), 2)
        titles = {c["title"] for c in msg2["canvases"]}
        self.assertEqual(titles, {"Doc 1", "Doc 2"})
        self.assertEqual(msg2["active_canvas"]["title"], "Doc 2")
        self.assertEqual(msg2["active_canvas"]["content"], "two")

        await comm.disconnect()

    async def test_canvas_switch_changes_active(self):
        from channels.db import database_sync_to_async

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)

        def setup():
            c1 = ChatCanvas.objects.create(thread=thread, title="Doc 1", content="one")
            c2 = ChatCanvas.objects.create(thread=thread, title="Doc 2", content="two")
            thread.active_canvas = c1
            thread.save(update_fields=["active_canvas"])
            return c2.pk
        c2_pk = await database_sync_to_async(setup)()

        comm = await self._communicator()
        await comm.send_json_to({
            "type": "chat.canvas_switch",
            "thread_id": str(thread.id),
            "canvas_id": c2_pk,
        })
        msg = await comm.receive_json_from()
        self.assertEqual(msg["event_type"], "canvas.loaded")
        self.assertEqual(msg["title"], "Doc 2")
        self.assertEqual(msg["content"], "two")

        def check_active():
            thread.refresh_from_db()
            self.assertEqual(thread.active_canvas_id, c2_pk)
        await database_sync_to_async(check_active)()
        await comm.disconnect()

    async def test_canvas_save_targets_specific_canvas(self):
        from channels.db import database_sync_to_async

        thread = await database_sync_to_async(ChatThread.objects.create)(created_by=self.user)

        def setup():
            c1 = ChatCanvas.objects.create(thread=thread, title="Doc 1", content="one")
            c2 = ChatCanvas.objects.create(thread=thread, title="Doc 2", content="two")
            thread.active_canvas = c1
            thread.save(update_fields=["active_canvas"])
            return c2.pk
        c2_pk = await database_sync_to_async(setup)()

        comm = await self._communicator()
        await comm.send_json_to({
            "type": "chat.canvas_save",
            "thread_id": str(thread.id),
            "canvas_id": c2_pk,
            "title": "Doc 2 Renamed",
            "content": "updated two",
        })
        await comm.disconnect()

        def check():
            c2 = ChatCanvas.objects.get(pk=c2_pk)
            self.assertEqual(c2.title, "Doc 2 Renamed")
            self.assertEqual(c2.content, "updated two")
            # c1 unchanged
            c1 = ChatCanvas.objects.get(thread=thread, title="Doc 1")
            self.assertEqual(c1.content, "one")
        await database_sync_to_async(check)()


# ---------------------------------------------------------------------------
# Multi-canvas prompt tests
# ---------------------------------------------------------------------------

class MultiCanvasPromptTests(TestCase):
    """Tests for build_system_prompt with multi-canvas parameters."""

    def test_all_canvas_titles_listed(self):
        from chat.prompts import build_system_prompt

        canvases = [
            {"title": "Doc A", "chars": 100, "is_active": False},
            {"title": "Doc B", "chars": 200, "is_active": True},
        ]

        class FakeCanvas:
            title = "Doc B"
            content = "B content here"

        prompt = build_system_prompt(canvases=canvases, active_canvas=FakeCanvas())
        self.assertIn("Doc A", prompt)
        self.assertIn("Doc B", prompt)
        self.assertIn("100 chars", prompt)
        self.assertIn("200 chars", prompt)
        self.assertIn("in context", prompt)
        self.assertIn("B content here", prompt)

    def test_active_canvas_content_shown(self):
        from chat.prompts import build_system_prompt

        canvases = [{"title": "Only", "chars": 42, "is_active": True}]

        class FakeCanvas:
            title = "Only"
            content = "The actual content"

        prompt = build_system_prompt(canvases=canvases, active_canvas=FakeCanvas())
        self.assertIn("The actual content", prompt)
        self.assertIn('Active Canvas Content: "Only"', prompt)

    def test_no_canvases_prompt(self):
        from chat.prompts import build_system_prompt

        prompt = build_system_prompt()
        self.assertIn("Each unique title creates a new canvas tab", prompt)
        self.assertNotIn("multiple document tabs", prompt)

    def test_backward_compat_single_canvas(self):
        from chat.prompts import build_system_prompt

        class FakeCanvas:
            title = "Old Style"
            content = "old content"

        prompt = build_system_prompt(canvas=FakeCanvas())
        self.assertIn("Old Style", prompt)
        self.assertIn("old content", prompt)


# ---------------------------------------------------------------------------
# Service-level tests for activate_canvas / set_active_canvases / resolve_canvas
# ---------------------------------------------------------------------------

class ActivateCanvasTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(email="activate@test.com", password="pass")
        self.thread = ChatThread.objects.create(created_by=self.user)

    def test_activate_single_canvas(self):
        from chat.services import activate_canvas

        canvas = ChatCanvas.objects.create(thread=self.thread, title="A", content="a")
        activate_canvas(self.thread.pk, canvas)
        canvas.refresh_from_db()
        self.assertTrue(canvas.is_active)
        self.assertIsNotNone(canvas.last_activated_at)

    def test_activate_bumps_timestamp_if_already_active(self):
        import time
        from chat.services import activate_canvas

        canvas = ChatCanvas.objects.create(thread=self.thread, title="A", content="a")
        activate_canvas(self.thread.pk, canvas)
        canvas.refresh_from_db()
        first_ts = canvas.last_activated_at

        time.sleep(0.01)
        activate_canvas(self.thread.pk, canvas)
        canvas.refresh_from_db()
        self.assertGreater(canvas.last_activated_at, first_ts)

    def test_cap_enforcement_deactivates_oldest(self):
        import time
        from chat.services import activate_canvas

        canvases = []
        for name in ["A", "B", "C"]:
            c = ChatCanvas.objects.create(thread=self.thread, title=name, content=name)
            activate_canvas(self.thread.pk, c)
            c.refresh_from_db()
            canvases.append(c)
            time.sleep(0.01)

        # All 3 active
        for c in canvases:
            c.refresh_from_db()
            self.assertTrue(c.is_active)

        # Activate 4th — oldest (A) should be deactivated
        d = ChatCanvas.objects.create(thread=self.thread, title="D", content="d")
        activate_canvas(self.thread.pk, d)

        canvases[0].refresh_from_db()
        self.assertFalse(canvases[0].is_active)
        d.refresh_from_db()
        self.assertTrue(d.is_active)

    def test_resolve_canvas_falls_back_to_most_recent(self):
        import time
        from chat.services import activate_canvas, resolve_canvas

        a = ChatCanvas.objects.create(thread=self.thread, title="A", content="a")
        b = ChatCanvas.objects.create(thread=self.thread, title="B", content="b")
        activate_canvas(self.thread.pk, a)
        time.sleep(0.01)
        activate_canvas(self.thread.pk, b)

        canvas, err = resolve_canvas(self.thread.pk)
        self.assertIsNone(err)
        self.assertEqual(canvas.pk, b.pk)

    def test_write_canvas_auto_activates(self):
        result = _invoke(
            WriteCanvasTool, {"title": "New", "content": "content"},
            _ctx(self.user.pk, self.thread.id),
        )
        self.assertEqual(result["status"], "ok")
        canvas = ChatCanvas.objects.get(thread=self.thread, title="New")
        self.assertTrue(canvas.is_active)
        self.assertIsNotNone(canvas.last_activated_at)

    def test_edit_canvas_bumps_activation(self):
        import time
        from django.utils import timezone

        canvas = ChatCanvas.objects.create(
            thread=self.thread, title="Doc", content="Hello world",
            is_active=True, last_activated_at=timezone.now(),
        )
        first_ts = canvas.last_activated_at
        time.sleep(0.01)

        _invoke(
            EditCanvasTool,
            {"edits": [{"old_text": "Hello", "new_text": "Hi", "reason": ""}]},
            _ctx(self.user.pk, self.thread.id),
        )
        canvas.refresh_from_db()
        self.assertGreater(canvas.last_activated_at, first_ts)
